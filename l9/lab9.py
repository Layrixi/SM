from io import BytesIO
import pathlib
import sys
import cv2
import numpy as np
import matplotlib.pyplot as plt
from tqdm import tqdm
from docx import Document
from docx.shared import Inches

##############################################################################
######   Konfiguracja       ##################################################
##############################################################################

kat = pathlib.Path().absolute() / 'l9'  # katalog z plikami wideo
plik="clip_2.mp4"                       # nazwa pliku
ile= 15                                 # ile klatek odtworzyć? <0 - całość
key_frame_counters = np.array([1,2,4,8])  # co ile klatek ma być kluczowa
key_frame_counter=4                     # co która klatka ma być kluczowa i nie podlegać kompresji
plot_frames=np.array([2,4,6,8,14])    # automatycznie wyrysuj wykresy
#plot_frames=np.array([1,2,3,6,10,14,18])    # automatycznie wyrysuj wykresy
auto_pause_frames=np.array([15])        # automatycznie za pauzuj dla klatki
subsampling="4:1:0"                     # parametry dla chroma subsampling
dzielnik=8                              # dzielnik przy zapisie różnicy
wyswietlaj_kaltki=False                 # czy program ma wyświetlać klatki
compressLuminance=False                  # czy kompresować luminancję    
ROI = [                                 # regiony obrazu do wyświetlenia  porównawczego     
        [100, 300, 100, 300], ]
loseless=False                          # czy ma być wykorzystany byterun   

##############################################################################
####     Kompresja i dekompresja    ##########################################
##############################################################################
class data:
    def init(self):
        self.Y=None
        self.Cb=None
        self.Cr=None
    def getSize(self):
        return get_size(self.Y)+get_size(self.Cb)+get_size(self.Cr)


def get_size(obj, seen=None):
    size = sys.getsizeof(obj)
    if seen is None:
        seen = set()
    obj_id = id(obj)
    if obj_id in seen:
        return 0
    # Important mark as seen *before* entering recursion to gracefully handle
    # self-referential objects
    seen.add(obj_id)
    if isinstance(obj,np.ndarray):
        size=obj.nbytes
    elif isinstance(obj, dict):
        size += sum([get_size(v, seen) for v in obj.values()])
        size += sum([get_size(k, seen) for k in obj.keys()])
    elif hasattr(obj, '__dict__'):
        size += get_size(obj.__dict__, seen)
    elif hasattr(obj, '__iter__') and not isinstance(obj, (str, bytes, bytearray)):
        size += sum([get_size(i, seen) for i in obj])
    return size


#zliczanie jak długo nie wystąpił ten sam symbol
def count_different_symbols(symbols, start):
    count =1
    for i in range(start+1, len(symbols)):
        if symbols[i] != symbols[i-1]:
            count += 1
        else:
            count-=1
            break
    return count
    

#zliczanie ile razy ten sam symbol
def count_repeated_symbols(symbols, start):
    count = 1
    while start + count < len(symbols) and symbols[start] == symbols[start + count]:
        count += 1
    return count


def ByteRun_compress(data):
    original_shape = data.shape
    data = data.flatten()  # Flatten the 2D array into 1D
    compressed = np.zeros(len(data) * 2, dtype=int)
    i = 0
    index = 0
    maxPackSize = 128

    while i < len(data):
        # If repetition
        if i + 1 < len(data) and data[i] == data[i + 1]:
            count = count_repeated_symbols(data, i)
            if count > maxPackSize:
                count = maxPackSize
            n = -count + 1
            compressed[index] = n
            compressed[index + 1] = data[i]
            index += 2
            i += count
        # If no repetition
        else:
            count = count_different_symbols(data, i)
            if count > maxPackSize:
                count = maxPackSize
            compressed[index] = count - 1
            for j in range(count):
                compressed[index + 1 + j] = data[i + j]
            index += count + 1
            i += count

    return compressed[:index], original_shape

def ByteRun_decompress(compressed_data, original_shape):
    decompressed = []
    i = 0
    compressed = compressed_data
    while i < len(compressed):
        n = compressed[i]
        if n < 0:
            # Repeated symbols
            count = -n + 1
            symbol = compressed[i + 1]
            decompressed.extend([symbol] * count)
            i += 2
        else:
            # Non-repeated symbols
            count = n + 1
            decompressed.extend(compressed[i + 1:i + 1 + count])
            i += 1 + count

    return np.array(decompressed, dtype=int).reshape(original_shape)

def Chroma_subsampling(L,subsampling):
     # L: chroma channel (2D numpy array)
    if subsampling == "4:4:4":
        return L.copy()
    elif subsampling == "4:2:2":
        # Horizontal subsampling by 2
        return L[:, ::2]
    elif subsampling == "4:2:0":
        # Horizontal and vertical subsampling by 2
        return L[::2, ::2]
    elif subsampling == "4:4:0":
        # Vertical subsampling by 2
        return L[::2, :]
    elif subsampling == "4:4:1":
        # Vertical subsampling by 4
        return L[::4, :]
    elif subsampling == "4:1:0":
        # Horizontal by 4, vertical by 2
        return L[::2, ::4]
    else:
        return L

def Chroma_resampling(L,subsampling):
    if subsampling == "4:4:4":
        return L.copy()
    elif subsampling == "4:2:2":
        # Repeat columns to restore original width
        return np.repeat(L, 2, axis=1)
    elif subsampling == "4:2:0":
        # Repeat rows and columns to restore original shape
        return np.repeat(np.repeat(L, 2, axis=0), 2, axis=1)
    elif subsampling == "4:4:0":
        # Repeat rows to restore original height
        return np.repeat(L, 2, axis=0)
    elif subsampling == "4:4:1":
        # Repeat rows to restore original height (by 4)
        return np.repeat(L, 4, axis=0)
    elif subsampling == "4:1:0":
        # Repeat rows by 2, columns by 4
        return np.repeat(np.repeat(L, 2, axis=0), 4, axis=1)
    else:
        return L

        
def frame_image_to_class(frame,subsampling):
    Frame_class = data()
    Frame_class.Y=frame[:,:,0].astype(int)
    Frame_class.Cb=Chroma_subsampling(frame[:,:,2].astype(int),subsampling)
    Frame_class.Cr=Chroma_subsampling(frame[:,:,1].astype(int),subsampling)
    return Frame_class


def frame_layers_to_image(Y,Cr,Cb,subsampling):  
    Cb=Chroma_resampling(Cb,subsampling)
    Cr=Chroma_resampling(Cr,subsampling)
    return np.dstack([Y,Cr,Cb]).clip(0,255).astype(np.uint8)

def compress_KeyFrame(Frame_class):
    KeyFrame = data()
    y_sub = Frame_class.Y
    cb_sub = Frame_class.Cb
    cr_sub = Frame_class.Cr
    if loseless:
        if compressLuminance:
            KeyFrame.Y, KeyFrame.Y_shape = ByteRun_compress(y_sub)
        else:
            KeyFrame.Y = y_sub
        KeyFrame.Cb, KeyFrame.Cb_shape = ByteRun_compress(cb_sub)
        KeyFrame.Cr, KeyFrame.Cr_shape = ByteRun_compress(cr_sub)
    else:
        KeyFrame.Y = y_sub
        KeyFrame.Cb = cb_sub
        KeyFrame.Cr = cr_sub
    return KeyFrame

def decompress_KeyFrame(KeyFrame):
    if loseless:
        if compressLuminance:
            Y = ByteRun_decompress(KeyFrame.Y, KeyFrame.Y_shape)
        else:
            Y = KeyFrame.Y
        Cb = ByteRun_decompress(KeyFrame.Cb, KeyFrame.Cb_shape)
        Cr = ByteRun_decompress(KeyFrame.Cr, KeyFrame.Cr_shape)
    else:
        Y = KeyFrame.Y
        Cb = KeyFrame.Cb
        Cr = KeyFrame.Cr
    frame_image = frame_layers_to_image(Y, Cr, Cb, subsampling)
    return frame_image

def compress_not_KeyFrame(Frame_class, KeyFrame, inne_paramerty_do_dopisania=None):
    Compress_data = data()
    if loseless:
        if compressLuminance:
            y_ref = ByteRun_decompress(KeyFrame.Y, KeyFrame.Y_shape)
            y_diff = (Frame_class.Y - y_ref) / dzielnik
            Compress_data.Y, Compress_data.Y_shape = ByteRun_compress(y_diff)
        else:
            y_ref = KeyFrame.Y
            y_diff = (Frame_class.Y - y_ref) / dzielnik
            Compress_data.Y = y_diff  # 2D, not compressed
        cb_ref = ByteRun_decompress(KeyFrame.Cb, KeyFrame.Cb_shape)
        cr_ref = ByteRun_decompress(KeyFrame.Cr, KeyFrame.Cr_shape)
        cb_diff = (Frame_class.Cb - cb_ref) / dzielnik
        cr_diff = (Frame_class.Cr - cr_ref) / dzielnik
        Compress_data.Cb, Compress_data.Cb_shape = ByteRun_compress(cb_diff)
        Compress_data.Cr, Compress_data.Cr_shape = ByteRun_compress(cr_diff)
    else:
        y_diff = (Frame_class.Y - KeyFrame.Y) / dzielnik
        cb_diff = (Frame_class.Cb - KeyFrame.Cb) / dzielnik
        cr_diff = (Frame_class.Cr - KeyFrame.Cr) / dzielnik
        Compress_data.Y = y_diff
        Compress_data.Cb = cb_diff
        Compress_data.Cr = cr_diff
    return Compress_data

def decompress_not_KeyFrame(Compress_data, KeyFrame, inne_paramerty_do_dopisania=None):
    if loseless:
        if compressLuminance:
            Y = (dzielnik * ByteRun_decompress(Compress_data.Y, Compress_data.Y_shape)) + ByteRun_decompress(KeyFrame.Y, KeyFrame.Y_shape)
        else:
            Y = (dzielnik * Compress_data.Y) + KeyFrame.Y  # Y is 2D, not compressed
        Cb = ByteRun_decompress(Compress_data.Cb, Compress_data.Cb_shape)
        Cr = ByteRun_decompress(Compress_data.Cr, Compress_data.Cr_shape)
        Cb = (dzielnik * Cb) + ByteRun_decompress(KeyFrame.Cb, KeyFrame.Cb_shape)
        Cr = (dzielnik * Cr) + ByteRun_decompress(KeyFrame.Cr, KeyFrame.Cr_shape)
    else:
        Y = (dzielnik * Compress_data.Y) + KeyFrame.Y
        Cb = (dzielnik * Compress_data.Cb) + KeyFrame.Cb
        Cr = (dzielnik * Compress_data.Cr) + KeyFrame.Cr
    return frame_layers_to_image(Y, Cr, Cb, subsampling)

def plotDiffrence(ReferenceFrame,DecompressedFrame,ROI,document=None):
    # bardzo słaby i sztuczny przykład wykorzystania tej opcji
    # przerobić żeby porównanie było dokonywane w RGB nie YCrCb i/lub zastąpić innym porównaniem
    # ROI - Region of Insert współrzędne fragmentu który chcemy przybliżyć i ocenić w formacie [w1,w2,k1,k2]
    # --- RGB ---
    fig, axs = plt.subplots(1, 3 , sharey=True   )
    fig.set_size_inches(16,5)
    ReferenceFrame_rgb = cv2.cvtColor(ReferenceFrame, cv2.COLOR_YCrCb2RGB)
    DecompressedFrame_rgb = cv2.cvtColor(DecompressedFrame, cv2.COLOR_YCrCb2RGB)
    axs[0].imshow(ReferenceFrame_rgb[ROI[0]:ROI[1],ROI[2]:ROI[3]])
    axs[2].imshow(DecompressedFrame_rgb[ROI[0]:ROI[1],ROI[2]:ROI[3]]) 
    diff=ReferenceFrame_rgb[ROI[0]:ROI[1],ROI[2]:ROI[3]].astype(float)-DecompressedFrame_rgb[ROI[0]:ROI[1],ROI[2]:ROI[3]].astype(float)
    print(np.min(diff),np.max(diff))
    axs[1].imshow(diff,vmin=np.min(diff),vmax=np.max(diff))
    #plot
    axs[0].set_title('Reference Frame')
    axs[1].set_title('Difference')
    axs[2].set_title('Decompressed Frame')
    axs[0].axis('off')
    axs[1].axis('off')
    axs[2].axis('off')
    plt.title("Porównanie klatek RGB")
    if document is not None:
        # Save the figure to a BytesIO object
        img_stream = BytesIO()
        plt.savefig(img_stream, format='png')
        img_stream.seek(0)
        # Add the image to the document
        document.add_picture(img_stream, width=Inches(6))
        # Close the BytesIO stream
        img_stream.close()
    else:
        plt.show()
    plt.close(fig)

    # --- YCrCb ---, just Y layer
    fig, axs = plt.subplots(1, 3 , sharey=True   )
    fig.set_size_inches(16,5)
    ReferenceFrame_y = ReferenceFrame[:, :, 0]
    DecompressedFrame_y = DecompressedFrame[:, :, 0]
    axs[0].imshow( (ReferenceFrame_y[ROI[0]:ROI[1],ROI[2]:ROI[3]]).clip(0, 255).astype(np.uint8), cmap='gray')
    axs[2].imshow( (DecompressedFrame_y[ROI[0]:ROI[1],ROI[2]:ROI[3]]).clip(0, 255).astype(np.uint8), cmap='gray')
    diff=ReferenceFrame_y[ROI[0]:ROI[1],ROI[2]:ROI[3]].astype(float)-DecompressedFrame_y[ROI[0]:ROI[1],ROI[2]:ROI[3]].astype(float)
    axs[1].imshow(diff,vmin=np.min(diff),vmax=np.max(diff), cmap='gray')
    axs[0].set_title('Reference Frame')
    axs[1].set_title('Difference')
    axs[2].set_title('Decompressed Frame')
    axs[0].axis('off')
    axs[1].axis('off')
    axs[2].axis('off')
    plt.title("Porównanie klatek warstwy Y")

    if document is not None:
        # Save the figure to a BytesIO object
        img_stream = BytesIO()
        plt.savefig(img_stream, format='png')
        img_stream.seek(0)
        # Add the image to the document
        document.add_picture(img_stream, width=Inches(6))
        # Close the BytesIO stream
        img_stream.close()
    else:
        plt.show()
    plt.close(fig)

    # --- Cb layer ---
    fig, axs = plt.subplots(1, 3 , sharey=True   )
    fig.set_size_inches(16,5)
    ReferenceFrame_y = ReferenceFrame[:, :, 1]
    DecompressedFrame_y = DecompressedFrame[:, :, 1]
    axs[0].imshow( (ReferenceFrame_y[ROI[0]:ROI[1],ROI[2]:ROI[3]]).clip(0, 255).astype(np.uint8), cmap='gray')
    axs[2].imshow( (DecompressedFrame_y[ROI[0]:ROI[1],ROI[2]:ROI[3]]).clip(0, 255).astype(np.uint8), cmap='gray')
    diff=ReferenceFrame_y[ROI[0]:ROI[1],ROI[2]:ROI[3]].astype(float)-DecompressedFrame_y[ROI[0]:ROI[1],ROI[2]:ROI[3]].astype(float)
    axs[1].imshow(diff,vmin=np.min(diff),vmax=np.max(diff), cmap='gray')
    axs[0].set_title('Reference Frame')
    axs[1].set_title('Difference')
    axs[2].set_title('Decompressed Frame')
    axs[0].axis('off')
    axs[1].axis('off')
    axs[2].axis('off')
    plt.title("Porównanie klatek Cr")

    if document is not None:
        # Save the figure to a BytesIO object
        img_stream = BytesIO()
        plt.savefig(img_stream, format='png')
        img_stream.seek(0)
        # Add the image to the document
        document.add_picture(img_stream, width=Inches(6))
        # Close the BytesIO stream
        img_stream.close()
    else:
        plt.show()
    plt.close(fig)

    #--- Cr layer ---
    fig, axs = plt.subplots(1, 3 , sharey=True   )
    fig.set_size_inches(16,5)
    ReferenceFrame_y = ReferenceFrame[:, :, 2]
    DecompressedFrame_y = DecompressedFrame[:, :, 2]
    axs[0].imshow( (ReferenceFrame_y[ROI[0]:ROI[1],ROI[2]:ROI[3]]).clip(0, 255).astype(np.uint8), cmap='gray')
    axs[2].imshow( (DecompressedFrame_y[ROI[0]:ROI[1],ROI[2]:ROI[3]]).clip(0, 255).astype(np.uint8), cmap='gray')
    diff=ReferenceFrame_y[ROI[0]:ROI[1],ROI[2]:ROI[3]].astype(float)-DecompressedFrame_y[ROI[0]:ROI[1],ROI[2]:ROI[3]].astype(float)
    axs[1].imshow(diff,vmin=np.min(diff),vmax=np.max(diff), cmap='gray')
    axs[0].set_title('Reference Frame')
    axs[1].set_title('Difference')
    axs[2].set_title('Decompressed Frame')
    axs[0].axis('off')
    axs[1].axis('off')
    axs[2].axis('off')
    plt.title("Porównanie klatek Cb")

    if document is not None:
        # Save the figure to a BytesIO object
        img_stream = BytesIO()
        plt.savefig(img_stream, format='png')
        img_stream.seek(0)
        # Add the image to the document
        document.add_picture(img_stream, width=Inches(6))
        # Close the BytesIO stream
        img_stream.close()
    else:
        plt.show()
    plt.close(fig)
    #--- YCrCb ---
    fig, axs = plt.subplots(1, 3 , sharey=True   )
    fig.set_size_inches(16,5)
    ReferenceFrame_y = ReferenceFrame
    DecompressedFrame_y = DecompressedFrame
    axs[0].imshow( (ReferenceFrame_y[ROI[0]:ROI[1],ROI[2]:ROI[3]]).clip(0, 255).astype(np.uint8), cmap='gray')
    axs[2].imshow( (DecompressedFrame_y[ROI[0]:ROI[1],ROI[2]:ROI[3]]).clip(0, 255).astype(np.uint8), cmap='gray')
    diff=ReferenceFrame_y[ROI[0]:ROI[1],ROI[2]:ROI[3]].astype(float)-DecompressedFrame_y[ROI[0]:ROI[1],ROI[2]:ROI[3]].astype(float)
    axs[1].imshow(diff,vmin=np.min(diff),vmax=np.max(diff), cmap='gray')
    axs[0].set_title('Reference Frame')
    axs[1].set_title('Difference')
    axs[2].set_title('Decompressed Frame')
    axs[0].axis('off')
    axs[1].axis('off')
    axs[2].axis('off')
    plt.title("Porównanie klatek YCrCb")

    if document is not None:
        # Save the figure to a BytesIO object
        img_stream = BytesIO()
        plt.savefig(img_stream, format='png')
        img_stream.seek(0)
        # Add the image to the document
        document.add_picture(img_stream, width=Inches(6))
        # Close the BytesIO stream
        img_stream.close()
    else:
        plt.show()
    plt.close(fig)




##############################################################################
####     Głowna pętla programu      ##########################################
##############################################################################
document = Document()
document.add_heading('Sprawozdanie 9',0) # tworzenie nagłówków, druga wartość to poziom nagłówka
#files = ["clip_2.mp4","clip_3.mp4"   ]
document.add_heading('Część 1 - badanie jakości dla różnych parametrów', level=1)
files = ["clip_2.mp4"] 
plik = files[0]
cap = cv2.VideoCapture(kat / plik)

if ile<0:
    ile=int(cap.get(cv2.CAP_PROP_FRAME_COUNT))

cv2.namedWindow('Normal Frame')
cv2.namedWindow('Decompressed Frame')

compression_information=np.zeros((3,ile))

document.add_heading('Plik: {}'.format(plik), level=1)
document.add_heading('Subsampling: {}'.format(subsampling), level=2)
document.add_heading('Dzielnik: {}'.format(dzielnik), level=2)
#####test part####
testowane_subsamplingi = ["4:4:4", "4:2:2", "4:2:0", "4:4:0", "4:1:0"]
testowane_key_frame_counters = key_frame_counters  

najlepsza_kompresja = -np.inf
najlepsze_ustawienia = (None, None)

for subs in tqdm(testowane_subsamplingi, desc="Testing subsampling"):
    for counter in tqdm(testowane_key_frame_counters, desc="Testing KeyFrame interval"):
        dz=counter
        subsampling = subs
        key_frame_counter = counter
        suma_poczatkowa = 0
        suma_po_kompresji = 0
        document.add_heading(f'Subsampling: {subsampling}, KeyFrame co: {key_frame_counter} klatki', level=2)
        for i in tqdm(range(ile), desc="Processing frames"):
            ret, frame = cap.read()
            if wyswietlaj_kaltki:
                cv2.imshow('Normal Frame', frame)
            frame = cv2.cvtColor(frame, cv2.COLOR_BGR2YCrCb)
            Frame_class = frame_image_to_class(frame, subsampling)
            original_size = get_size(frame)

            #######kompresja i dekompresja klatek
            if (i % key_frame_counter) == 0:
                KeyFrame = compress_KeyFrame(Frame_class)
                compressed_size = get_size(KeyFrame)
                cY = KeyFrame.Y
                cCb = KeyFrame.Cb
                cCr = KeyFrame.Cr
                Decompresed_Frame = decompress_KeyFrame(KeyFrame)
            else:
                Compress_data = compress_not_KeyFrame(Frame_class, KeyFrame)
                compressed_size = get_size(Compress_data)
                cY = Compress_data.Y
                cCb = Compress_data.Cb
                cCr = Compress_data.Cr
                Decompresed_Frame = decompress_not_KeyFrame(Compress_data, KeyFrame)

            compression_information[0, i] = (frame[:, :, 0].size - cY.size) / frame[:, :, 0].size
            compression_information[1, i] = (frame[:, :, 0].size - cCb.size) / frame[:, :, 0].size
            compression_information[2, i] = (frame[:, :, 0].size - cCr.size) / frame[:, :, 0].size

            ######sprawozdanie
            if wyswietlaj_kaltki:
                cv2.imshow('Decompressed Frame', cv2.cvtColor(Decompresed_Frame, cv2.COLOR_YCrCb2BGR))

            if np.any(plot_frames == i):
                for r in ROI:
                    if document is not None:
                        document.add_paragraph(f'Region: {r}, klatka: {i}')
                        plotDiffrence(frame, Decompresed_Frame, r, document)
                    else:
                        plotDiffrence(frame, Decompresed_Frame, r)

            if np.any(auto_pause_frames == i):
                cv2.waitKey(-1) #wait until any key is pressed
            
            k = cv2.waitKey(1) & 0xff
            
            if k==ord('q'):
                break
            elif k == ord('p'):
                cv2.waitKey(-1) #wait until any key is pressed
            
            suma_poczatkowa += original_size
            suma_po_kompresji += compressed_size
        cap.set(cv2.CAP_PROP_POS_FRAMES, 0)
        stopien_kompresji = suma_poczatkowa / suma_po_kompresji
        print(f"Test: {subs}, dzielnik={dz}, CR={stopien_kompresji:.2f}")
        document.add_paragraph(f'Test: {subs}, dzielnik={dz}, Stopień kompresji: {stopien_kompresji:.2f}')
        if stopien_kompresji > najlepsza_kompresja:
            najlepsza_kompresja = stopien_kompresji
            najlepsze_ustawienia = (subs, dz)
    print(f"\n>> Najlepsze ustawienia: {subsampling}, dzielnik={dzielnik}, kompresja={najlepsza_kompresja:.2f}")
document.add_heading('Najlepsze ustawienia: {}, dzielnik={}, kompresja={:.2f}'.format(najlepsze_ustawienia[0], najlepsze_ustawienia[1], najlepsza_kompresja), level=2)

document.add_heading('Część 2 - kompresja bezstratna', level=1)

subsampling = najlepsze_ustawienia[0]  # ustawienie najlepszych parametrów
key_frame_counter = najlepsze_ustawienia[1]  # ustawienie najlepszych parametrów
#cz2

#########################################################
########################################################
######################################################
loseless = True  # ustawienie kompresji bezstratnej


compressLuminance = True
key_frame_counters = np.array([1,2,3,4,7,8])  # co ile klatek ma być kluczowa

cap.set(cv2.CAP_PROP_POS_FRAMES, 0)
compression_information = np.zeros((3, ile))

if document is not None:
    document.add_heading('Kompresja z luminancją', level=2)
for key_frame_counter in key_frame_counters:
    print(key_frame_counter)
    dzielnik = key_frame_counter
    cap.set(cv2.CAP_PROP_POS_FRAMES, 0)  # resetowanie pozycji odtwarzania
    for i in tqdm(range(ile), desc="Processing frames (with luminance compression)"):
        ret, frame = cap.read()
        if not ret:
            break

        if wyswietlaj_kaltki:
            cv2.imshow('Normal Frame', frame)

        frame = cv2.cvtColor(frame, cv2.COLOR_BGR2YCrCb)
        Frame_class = frame_image_to_class(frame, subsampling)
        original_size = get_size(frame)

        if (i % key_frame_counter) == 0:
            KeyFrame = compress_KeyFrame(Frame_class)
            compressed_size = get_size(KeyFrame)
            cY = KeyFrame.Y
            cCb = KeyFrame.Cb
            cCr = KeyFrame.Cr
            Decompresed_Frame = decompress_KeyFrame(KeyFrame)
        else:
            Compress_data = compress_not_KeyFrame(Frame_class, KeyFrame)
            compressed_size = get_size(Compress_data)
            cY = Compress_data.Y
            cCb = Compress_data.Cb
            cCr = Compress_data.Cr
            Decompresed_Frame = decompress_not_KeyFrame(Compress_data, KeyFrame)

        compression_information[0, i] = (frame[:, :, 0].size - cY.size) / frame[:, :, 0].size
        compression_information[1, i] = (frame[:, :, 0].size - cCb.size) / frame[:, :, 0].size
        compression_information[2, i] = (frame[:, :, 0].size - cCr.size) / frame[:, :, 0].size

        if wyswietlaj_kaltki:
            cv2.imshow('Decompressed Frame', cv2.cvtColor(Decompresed_Frame, cv2.COLOR_YCrCb2BGR))


        if np.any(auto_pause_frames == i):
            cv2.waitKey(-1)

        k = cv2.waitKey(1) & 0xff
        if k == ord('q'):
            break
        elif k == ord('p'):
            cv2.waitKey(-1)

    # wykres kompresji
    plt.figure()
    plt.plot(np.arange(0, ile), compression_information[0, :] * 100, label='Y')
    plt.plot(np.arange(0, ile), compression_information[1, :] * 100, label='Cb')
    plt.plot(np.arange(0, ile), compression_information[2, :] * 100, label='Cr')
    plt.title("File:{}, subsampling={}, divider={}, KeyFrame={},\n Kompresja bezstratna ByteRun".format(plik, subsampling, dzielnik, key_frame_counter))
    plt.xlabel('Frame')
    plt.ylabel('Compression [%]')
    plt.legend()

    img_stream = BytesIO()
    plt.savefig(img_stream, format='png')
    img_stream.seek(0)
    if document is not None:
        document.add_picture(img_stream, width=Inches(6))
    img_stream.close()
    plt.close()

#bez kompresji luminancji
compressLuminance = False
cap.set(cv2.CAP_PROP_POS_FRAMES, 0)
compression_information = np.zeros((3, ile))

if document is not None:
    document.add_heading('Kompresja bez luminancji', level=2)
for key_frame_counter in key_frame_counters:
    print(key_frame_counter)
    dzielnik = key_frame_counter
    cap.set(cv2.CAP_PROP_POS_FRAMES, 0)  # resetowanie pozycji odtwarzania
    for i in tqdm(range(ile), desc="Processing frames (with luminance compression)"):
        ret, frame = cap.read()
        if not ret:
            break

        if wyswietlaj_kaltki:
            cv2.imshow('Normal Frame', frame)

        frame = cv2.cvtColor(frame, cv2.COLOR_BGR2YCrCb)
        Frame_class = frame_image_to_class(frame, subsampling)
        original_size = get_size(frame)

        if (i % key_frame_counter) == 0:
            KeyFrame = compress_KeyFrame(Frame_class)
            compressed_size = get_size(KeyFrame)
            cY = KeyFrame.Y
            cCb = KeyFrame.Cb
            cCr = KeyFrame.Cr
            Decompresed_Frame = decompress_KeyFrame(KeyFrame)
        else:
            Compress_data = compress_not_KeyFrame(Frame_class, KeyFrame)
            compressed_size = get_size(Compress_data)
            cY = Compress_data.Y
            cCb = Compress_data.Cb
            cCr = Compress_data.Cr
            Decompresed_Frame = decompress_not_KeyFrame(Compress_data, KeyFrame)

        compression_information[0, i] = (frame[:, :, 0].size - cY.size) / frame[:, :, 0].size
        compression_information[1, i] = (frame[:, :, 0].size - cCb.size) / frame[:, :, 0].size
        compression_information[2, i] = (frame[:, :, 0].size - cCr.size) / frame[:, :, 0].size

        if wyswietlaj_kaltki:
            cv2.imshow('Decompressed Frame', cv2.cvtColor(Decompresed_Frame, cv2.COLOR_YCrCb2BGR))

       

        if np.any(auto_pause_frames == i):
            cv2.waitKey(-1)

        k = cv2.waitKey(1) & 0xff
        if k == ord('q'):
            break
        elif k == ord('p'):
            cv2.waitKey(-1)

    # wykres kompresji
    plt.figure()
    plt.plot(np.arange(0, ile), compression_information[0, :] * 100, label='Y')
    plt.plot(np.arange(0, ile), compression_information[1, :] * 100, label='Cb')
    plt.plot(np.arange(0, ile), compression_information[2, :] * 100, label='Cr')
    plt.title("File:{}, subsampling={}, divider={}, KeyFrame={},\n Kompresja bezstratna: ByteRun".format(plik, subsampling, dzielnik, key_frame_counter))
    plt.xlabel('Frame')
    plt.ylabel('Compression [%]')
    plt.legend()

    img_stream = BytesIO()
    plt.savefig(img_stream, format='png')
    img_stream.seek(0)
    if document is not None:
        document.add_picture(img_stream, width=Inches(6))
    img_stream.close()
    plt.close()

    
# dla innego filmu
document.add_heading('Część 3 - inny film', level=1)
plik = "clip_3.mp4"
document.add_heading('Plik: {}'.format(plik), level=2)
cap = cv2.VideoCapture(kat / plik)
compressLuminance = True  # ustawienie kompresji bezstratnej
key_frame_counter = 3  # co która klatka ma być kluczowa i nie podlegać kompresji

dzielnik = key_frame_counter
cap.set(cv2.CAP_PROP_POS_FRAMES, 0)  # resetowanie pozycji odtwarzania
for i in tqdm(range(ile), desc="Processing frames (with luminance compression)"):
    ret, frame = cap.read()
    if not ret:
        break

    if wyswietlaj_kaltki:
        cv2.imshow('Normal Frame', frame)

    frame = cv2.cvtColor(frame, cv2.COLOR_BGR2YCrCb)
    Frame_class = frame_image_to_class(frame, subsampling)
    original_size = get_size(frame)

    if (i % key_frame_counter) == 0:
        KeyFrame = compress_KeyFrame(Frame_class)
        compressed_size = get_size(KeyFrame)
        cY = KeyFrame.Y
        cCb = KeyFrame.Cb
        cCr = KeyFrame.Cr
        Decompresed_Frame = decompress_KeyFrame(KeyFrame)
    else:
        Compress_data = compress_not_KeyFrame(Frame_class, KeyFrame)
        compressed_size = get_size(Compress_data)
        cY = Compress_data.Y
        cCb = Compress_data.Cb
        cCr = Compress_data.Cr
        Decompresed_Frame = decompress_not_KeyFrame(Compress_data, KeyFrame)

    compression_information[0, i] = (frame[:, :, 0].size - cY.size) / frame[:, :, 0].size
    compression_information[1, i] = (frame[:, :, 0].size - cCb.size) / frame[:, :, 0].size
    compression_information[2, i] = (frame[:, :, 0].size - cCr.size) / frame[:, :, 0].size

    if wyswietlaj_kaltki:
        cv2.imshow('Decompressed Frame', cv2.cvtColor(Decompresed_Frame, cv2.COLOR_YCrCb2BGR))

    

    if np.any(auto_pause_frames == i):
        cv2.waitKey(-1)

    k = cv2.waitKey(1) & 0xff
    if k == ord('q'):
        break
    elif k == ord('p'):
        cv2.waitKey(-1)

# wykres kompresji
plt.figure()
plt.plot(np.arange(0, ile), compression_information[0, :] * 100, label='Y')
plt.plot(np.arange(0, ile), compression_information[1, :] * 100, label='Cb')
plt.plot(np.arange(0, ile), compression_information[2, :] * 100, label='Cr')
plt.title("File:{}, subsampling={}, divider={}, KeyFrame={},\n Kompresja bezstratna: ByteRun".format(plik, subsampling, dzielnik, key_frame_counter))
plt.xlabel('Frame')
plt.ylabel('Compression [%]')
plt.legend()

img_stream = BytesIO()
plt.savefig(img_stream, format='png')
img_stream.seek(0)
if document is not None:
    document.add_picture(img_stream, width=Inches(6))
img_stream.close()
plt.close()

document.save('sprawozdanie9.docx')




