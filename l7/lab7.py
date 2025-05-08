import numpy as np
import matplotlib.pyplot as plt
import scipy.fftpack
import sounddevice as sd
import soundfile as sf
import pathlib
from docx import Document
from docx.shared import Inches
from io import BytesIO
from scipy.interpolate import interp1d

def PlotToDoc(document,fig):
    fig.tight_layout(pad=1.5) # poprawa czytelności 
    memfile = BytesIO() # tworzenie bufora
    fig.savefig(memfile) # z zapis do bufora
    document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku
    memfile.close()
def plotSubplot2(x1,y1,x2,y2,title1="plot1",title2="plot2",document=None):
        fig ,axs = plt.subplots(2,1,figsize=(10,7)) # tworzenie plota
        plt.subplot(2,1,1)
        plt.xlim(np.min(x1),np.max(x1))
        plt.ylim(np.min(y1),np.max(y1))
        plt.xlabel("X")
        plt.ylabel("Y")
        plt.plot(x1,y1)

        plt.subplot(2,1,2)
        plt.xlim(np.min(x2),np.max(x2))
        plt.ylim(np.min(y2),np.max(y2))
        plt.xlabel("X")
        plt.ylabel("Y")
        plt.plot(x2,y2)
        if document is not None:
            fig.tight_layout(pad=1.5) # poprawa czytelności 
            memfile = BytesIO() # tworzenie bufora
            fig.savefig(memfile) # z zapis do bufora
            document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku
            memfile.close()
        else:
            plt.show()
        plt.close()

def plotSubplot3(x1,y1,x2,y2,x3,y3,title1="plot1",title2="plot2",title3="plot3",document=None):
        fig ,axs = plt.subplots(1,3,figsize=(10,7)) # tworzenie plota
        plt.subplot(3,1,1)
        plt.xlim(np.min(x1),np.max(x1))
        plt.ylim(np.min(y1),np.max(y1))
        plt.xlabel("X")
        plt.ylabel("Y")
        plt.title(title1)
        plt.plot(x1,y1)

        plt.subplot(3,1,2)
        plt.xlim(np.min(x2),np.max(x2))
        plt.ylim(np.min(y2),np.max(y2))
        plt.xlabel("X")
        plt.ylabel("Y")
        plt.title(title2)
        plt.plot(x2,y2)

        plt.subplot(3,1,3)
        plt.xlim(np.min(x3),np.max(x3))
        plt.ylim(np.min(y3),np.max(y3))
        plt.xlabel("X")
        plt.ylabel("Y")
        plt.title(title3)
        plt.plot(x3,y3)

        if document is not None:
            fig.tight_layout(pad=1.5) # poprawa czytelności 
            memfile = BytesIO() # tworzenie bufora
            fig.savefig(memfile) # z zapis do bufora
            document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku
            memfile.close()
        else:
            plt.show()
        plt.close()

def plotAudio(signal,fs,fsize,document=None,TimeMargin=[0,0.02],title="Audio"):
        fig ,axs = plt.subplots(2,1,figsize=(10,7)) # tworzenie plota
        fig.suptitle(title) 
        plt.subplot(2,1,1)
        plt.xlim(TimeMargin)
        plt.ylim(np.min(signal), np.max(signal)) 
        plt.xlabel("T(s)")
        plt.ylabel("Amplitude")
        plt.plot(np.arange(0,signal.shape[0])/fs,signal)


        plt.subplot(2,1,2)
        yf = scipy.fftpack.fft(signal,fsize)
        xf=np.arange(0,fs/2,fs/fsize)
        dbspectrum = 20 * np.log10(np.abs(yf[:fsize // 2])+np.finfo(np.float64).eps)
        plt.xlabel("Fs(Hz)")
        plt.ylabel("Magnitude(dB)")
        plt.plot(np.arange(0,fs/2,fs/fsize),dbspectrum)
        max_magnitude = np.max(dbspectrum)
        max_frequency = xf[np.argmax(dbspectrum)]
        if document is not None:
                fig.tight_layout(pad=1.5) # poprawa czytelności 
                memfile = BytesIO() # tworzenie bufora
                fig.savefig(memfile) # z zapis do bufora
                document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku
                document.add_paragraph('Max db: {} dB w fs: {} Hz'.format(max_magnitude,max_frequency))
                memfile.close()
        else:
                plt.show()
        plt.close()

def kwant(data, bit):
        d = 2**bit - 1
        
        if np.issubdtype(data.dtype, np.floating):
                m = -1
                n = 1
        else:
                m = np.iinfo(data.dtype).min
                n = np.iinfo(data.dtype).max

        #1
        DataF = data.astype(float)
        DataF = (DataF - m) / (n - m)

        #2
        DataF = np.round(DataF * d)
        DataF = DataF / d

        #3
        DataF = DataF * (n - m) + m

        return DataF.astype(data.dtype)

#DPCM Z PREDYKCJĄ NIE DZIAŁA POPRAWNIE.
#dpcm bez predykcji
def DPCM_compress(x,bit):
    y=np.zeros(x.shape)
    e=0
    for i in range(0,x.shape[0]):
        y[i]=kwant((x[i]-e).astype(x.dtype),bit)
        e+=y[i]
    return y

#xp jrst dekompresja
def DPCM_decompress(y):
    x=np.zeros(y.shape)
    x[0]=y[0]
    for i in range(1,y.shape[0]):
        x[i]=y[i]+x[i-1]
    return x
#dpcm z predykcją, najprostsza to np.mean
def DPCM_Pred_Compress(x,bit,predictor,n,ceil=False): 
    y=np.zeros(x.shape)
    xp=np.zeros(x.shape)
    e=0
    for i in range(0,x.shape[0]):
        y[i]=kwant(np.floor(x[i]-e).astype(x.dtype),bit)
        xp[i]=y[i]+e
        idx=(np.arange(i-n,i,1,dtype=int)+1)
        idx=np.delete(idx,idx<0)
        e=predictor(xp[idx],ceil=ceil)
    return y

def DPCM_Pred_Decompress(y, predictor, n, ceil=False):
    x = np.zeros(y.shape)
    e = 0
    for i in range(0, y.shape[0]):
        x[i] = np.floor(y[i] + e).astype(y.dtype)
        idx = (np.arange(i - n, i, 1, dtype=int) + 1)
        idx = np.delete(idx, idx < 0)
        e = predictor(x[idx], ceil=ceil)
    return x
    
def no_pred(X):
    return X[-1]

def mean_pred(X, ceil=True):
    if ceil:
        return np.trunc(np.mean(X))
    else:
        return np.mean(X)

def median_pred(X,ceil=True):
    if ceil:
        return np.ceil(np.median(X))
    else:
        return np.median(X)
def ALawCompress(x):
    A=87.6
    y=np.zeros(x.shape)
    idx=np.abs(x)<(1/A)
    y[idx]=np.sign(x[idx]) * ( A*np.abs(x[idx]) / (1+np.log(A)))
    idx=((1/A) <= np.abs(x)) & (np.abs(x) <= 1)
    y[idx]=np.sign(x[idx]) * ( (1+np.log(A*np.abs(x[idx]))) / (1+np.log(A)) ) 
    return y

def ALawDecompress(y):
    A=87.6
    x=np.zeros(y.shape)
    idy=np.abs(y) < (1/(1+np.log(A)))
    x[idy]=np.sign(y[idy]) * (np.abs(y[idy])*(1+np.log(A))/A)
    idy=(1/(1+np.log(A)) <= np.abs(y)) & (np.abs(y) <= 1)
    x[idy]=np.sign(y[idy]) * ( (np.exp( np.abs(y[idy]) * ( 1 + np.log(A)) - 1)) / A )
    return x

def MuLawCompress(x):
    mu=255
    y=np.zeros(x.shape)
    idx=(np.abs(x)<= 1 ) & (np.abs(x)>=-1)
    y[idx]=np.sign(x[idx]) * ( np.log( 1 + (mu * np.abs(x[idx]) ) ) / np.log( 1 + mu )  )
    return y

def MuLawDecompress(y):
    mu=255
    x=np.zeros(y.shape)
    idy=np.abs(y)<=(1)
    x[idy] = np.sign(y[idy]) * (1 / mu) * (np.power(1 + mu, np.abs(y[idy])) - 1)
    return y


"""#mulaw and alaw testing
x=np.linspace(-1,1,1000)
y=0.9*np.sin(np.pi*x*4)
plt.plot(x)
plt.show()
xALawCompressed=ALawCompress(x)
plt.plot(x,xALawCompressed)
plt.show()
xALawyCompressedKwant = kwant(xALawCompressed,6)

xALawDecompressed=ALawDecompress(xALawyCompressedKwant)
plt.plot(xALawDecompressed)
plt.show()

xMuLawCompressed=MuLawCompress(x)
plt.plot(x,xMuLawCompressed)
plt.show()
xMuLawCompressedKwant = kwant(xMuLawCompressed,6)
plt.plot(x,xMuLawCompressedKwant)
plt.show()
#"""

"""#logical indexing test
R=np.random.rand(5,5)
A=np.zeros(R.shape)
B=np.zeros(R.shape)
C=np.zeros(R.shape)


idx=R<0.25
A[idx]=1 # <-
B[idx]+=0.25 # <-
C[idx]=2*R[idx]-0.25 # <-
C[np.logical_not(idx)]=4*R[np.logical_not(idx)]-0.5 # <-
print(R)
print(idx)
print(A)
print(B)
print(C)#"""

"""#tests v2
x=np.linspace(-1,1,1000)
y=0.9*np.sin(np.pi*x*4)
plt.plot(x,y)
plt.show()

sinALawCompressed=kwant(ALawCompress(y),6)

sinMuLawCompressed=kwant(MuLawCompress(y),6)

plotSubplot2(x,sinALawCompressed,x,sinMuLawCompressed,"sinALaw","sinMuLaw")#"""
"""
#DPCM test
#x=np.array([15,16,20,14,5,10,15,13,11,7,10,11,20,1,23],dtype=np.int8)
x=np.linspace(-1,1,1000)
y=0.9*np.sin(np.pi*x*4)
xCompressed=DPCM_compress(y,7)
xDecompressed=DPCM_decompress(xCompressed)
plt.plot(x,xDecompressed)
plt.show()
print("x: \n",x)
print("xCompressed: \n",xCompressed)
print("xDecompressed:\n",xDecompressed)#"""

""""""
# Przygotowanie dokumentu
document = Document()
document.add_heading('Raport 7 Jacewicz Jacek', level=1)

#ALaw x in x out
x = np.linspace(-1, 1, 1000)

compressed = DPCM_compress(x,7)
decompressed = DPCM_decompress(compressed)
document.add_heading('ALaw i MuLaw test na przestrzeni sygnału', level=2)
# Test signal

xALawCompressed = ALawCompress(x)
xALawyCompressedKwant = kwant(xALawCompressed, 8)
xMuLawCompressed = MuLawCompress(x)
xMuLawCompressedKwant = kwant(xMuLawCompressed, 8)

# Plotting all 4 on one subplot
plt.figure(figsize=(10, 7))
plt.plot(x, xALawCompressed, label="ALaw")
plt.plot(x, xALawyCompressedKwant, label="ALaw + Kwantyzacja(8bit)")
plt.plot(x, xMuLawCompressed, label="MuLaw")
plt.plot(x, xMuLawCompressedKwant, label="MuLaw + Kwantyzacja(8bit)")
plt.xlabel("Wartość sygnału wejsciowego")
plt.ylabel("Wartość sygnału wyjsciowego")
plt.title("Krzywa po kompresji")
plt.legend()
plt.grid()
PlotToDoc(document,plt)
plt.close()

xMuLawDecompressed=MuLawDecompress(xMuLawCompressedKwant)
xALawDecompressed=ALawDecompress(xALawyCompressedKwant)
xKwant =kwant(x,8)
plt.figure(figsize=(10, 7))
plt.plot(x,x, label="Sygnał oryginalny")
plt.plot(x, xALawDecompressed, label="Sygnał po dekompresji z ALaw(8 bit)")
plt.plot(x, xMuLawDecompressed, label="Sygnał po dekompresji z MuLaw(8bit)")
plt.plot(x, xKwant, label="Sygnał oryginalny po kwantyzacji(8bit)")
plt.xlabel("Wartość sygnału wejsciowego")
plt.ylabel("Wartość sygnału wyjsciowego")
plt.title("Krzywa po dekompresji")
plt.legend()
plt.grid()
PlotToDoc(document,plt)
plt.close()

# dpcm
x=np.array([15,16,20,14,5,10,15,13,11,7,10,11,20,1,23],dtype=np.int8)
xCompressed=DPCM_compress(x,7)
xDecompressed=DPCM_decompress(xCompressed)
print("x: \n",x)
print("xCompressed: \n",xCompressed)
print("xDecompressed:\n",xDecompressed)#
document.add_heading('DPCM test na x=np.array([15,16,20,14,5,10,15,13,11,7,10,11,20,1,23]), kwantyzacja do 7 bitów', level=2)
document.add_paragraph(f'x:\n {x}')
document.add_paragraph(f'xCompressed: \n{xCompressed}')
document.add_paragraph(f'xDecompressed: \n{xDecompressed}')

# dpcm pred
xCompressed=DPCM_Pred_Compress(x,7,mean_pred,3,ceil=True)
xDecompressed=DPCM_Pred_Decompress(xCompressed,mean_pred,3,ceil=True)
print("x: \n",x)
print("xCompressed: \n",xCompressed)
print("xDecompressed:\n",xDecompressed)#
document.add_heading('DPCM z predykcją, metoda predykcji: średnia 3 ostatnich wartości z zaokrąglaniem wartości, test na x=np.array([15,16,20,14,5,10,15,13,11,7,10,11,20,1,23]), kwantyzacja do 7 bitów', level=2)
document.add_paragraph(f'x:\n {x}')
document.add_paragraph(f'xCompressed: \n{xCompressed}')
document.add_paragraph(f'xDecompressed: \n{xDecompressed}')

document.add_heading('Test na sygnale sinusoidalnym, sygnały wyjściowe po dekompresji', level=2)
#sin graphs
x=np.linspace(-1,1,1000)
y=0.9*np.sin(np.pi*x*4)
plt.figure(figsize=(10, 7))

plt.subplot(5,1,1)
plt.plot(x,y)
plt.title("Sygnał oryginalny")

plt.subplot(5,1,2)
xCompressed=ALawCompress(y)
xCompressed=kwant(xCompressed,7)
xDecompressed=ALawDecompress(xCompressed)
plt.plot(x,xDecompressed)
plt.title("Kompresja ALaw")

plt.subplot(5,1,3)
xCompressed=MuLawCompress(y)
xCompressed=kwant(xCompressed,7)
xDecompressed=MuLawDecompress(xCompressed)
plt.plot(x,xDecompressed)
plt.title("Kompresja MuLaw")

plt.subplot(5,1,4)
xCompressed=DPCM_compress(y,7)
xDecompressed=DPCM_decompress(xCompressed)
plt.plot(x,xDecompressed)
plt.title("Kompresja DPCM")

plt.subplot(5,1,5)
xCompressed=DPCM_Pred_Compress(y,7,mean_pred,3,ceil=True)
xDecompressed=DPCM_Pred_Decompress(xCompressed,mean_pred,3,ceil=True)
plt.plot(x,xDecompressed)
plt.title("Kompresja DPCM z predykcją")
plt.tight_layout()
PlotToDoc(document,plt)
plt.close()


singPath = pathlib.Path().absolute() / 'l7' / 'SING' 
fsize=2**8
singFiles=['sing_high1.wav','sing_low1.wav','sing_medium1.wav']
bits_list = [8,7,6,5,4,3,2]
for file in singFiles:
    # Wczytanie pliku
    data, fs = sf.read(singPath / file, dtype=np.int32)

    for bits in bits_list:
        # ALaw Compression and Decompression
        compressed = kwant(ALawCompress(data), bits)
        decompressed = ALawDecompress(compressed)
        sf.write(singPath / f"{file}_ALaw_{bits}bit.wav", decompressed, fs)
        
        # MuLaw Compression and Decompression
        compressed = kwant(MuLawCompress(data), bits)
        decompressed = MuLawDecompress(compressed)
        sf.write(singPath / f"{file}_MuLaw_{bits}bit.wav", decompressed, fs)

        # DPCM Compression and Decompression
        compressed = DPCM_compress(data, bits)
        decompressed = DPCM_decompress(compressed)
        sf.write(singPath / f"{file}_DPCM_{bits}bit.wav", decompressed, fs)

        # DPCM with Prediction Compression and Decompression
        compressed = DPCM_Pred_Compress(data, bits, mean_pred, 3, ceil=True)
        decompressed = DPCM_Pred_Decompress(compressed, mean_pred, 3, ceil=True)
        sf.write(singPath / f"{file}_DPCM_Pred_{bits}bit.wav", decompressed, fs)
# Zapis dokumentu
savepath = pathlib.Path().absolute() / 'results.docx'
document.save(savepath)#"""
