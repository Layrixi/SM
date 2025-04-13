import numpy as np
import matplotlib.pyplot as plt
import cv2
import pathlib
from docx import Document
from docx.shared import Inches
from io import BytesIO
from os import path
from tqdm import tqdm


def colorFit(pixel,Pallet):
        dist = np.linalg.norm(Pallet-pixel,axis=1)
        return Pallet[np.argmin(dist)]

    
def kwant_colorFit(img,Pallet):
        out_img = img.copy()
        for w in tqdm(range(img.shape[0]), desc="Processing rows"):
                for k in range(img.shape[1]):
                        out_img[w,k]=colorFit(img[w,k],Pallet)
        return out_img

def randDithering(img):
        r=np.random.rand(img.shape[0],img.shape[1])
        out_img = img >= r
        out_img = out_img * 1
        return out_img

def organizedDithering(img,Pallet,n=2,r=1,M=np.array([[0, 8, 2, 10],[12, 4, 14, 6],[3, 11, 1, 9],[15, 7, 13, 5]])):
        out_img = img.copy()
        Mpre = (M+1) / (2*n)**2 - 0.5
        height, width = img.shape[0], img.shape[1]

        for i in tqdm(range(height), desc="Processing rows"):
                for j in range(width):
                        out_img[i, j] = colorFit(out_img[i,j]+r*Mpre[i%(2*n),j%(2*n)],Pallet)
        print("")
        return out_img

def FloydDithering(img,Pallet):
        out_img = img.copy()
        height, width = img.shape[0], img.shape[1]
        
        for i in tqdm(range(height), desc="Processing rows"):
                for j in range(width):
                        oldpixel = out_img[i,j].copy()
                        newpixel = colorFit(oldpixel,Pallet)
                        out_img[i,j] = newpixel
                        quant_error = np.array(oldpixel - newpixel)
                        if i + 1 < height:
                                out_img[i + 1, j] = out_img[i + 1, j] + quant_error * 7 / 16
                        if i - 1 >= 0 and j + 1 < width:
                            out_img[i - 1, j + 1] = out_img[i - 1, j + 1] + quant_error * 3 / 16
                        if j + 1 < width:
                            out_img[i, j + 1] = out_img[i, j + 1] + quant_error * 5 / 16
                        if i + 1 < height and j + 1 < width:
                            out_img[i + 1, j + 1] = out_img[i + 1, j + 1] + quant_error * 1 / 16
        print("")
        return out_img
def imgToUInt8(img):
    if np.issubdtype(img.dtype,np.float32):
        img = (img*255).astype(np.uint8)
    return img
def imgToFloat(img):
    if np.issubdtype(img.dtype,np.uint8):
        img = img.astype(np.float32)/255.
    return img

pathSmall = pathlib.Path().absolute() /'l4'/'IMG_SMALL'
pathGs = pathlib.Path().absolute() /'l4'/'IMG_GS'
filesSmall=['SMALL_0001.tif',
        'SMALL_0004.jpg',
        'SMALL_0006.jpg',
        'SMALL_0007.jpg']
filesGs=['GS_0001.tif',
         'GS_0002.png',
         'GS_0003.png']
pallet16 = np.array([
        [0.0, 0.0, 0.0,], 
        [0.0, 1.0, 1.0,],
        [0.0, 0.0, 1.0,],
        [1.0, 0.0, 1.0,],
        [0.0, 0.5, 0.0,], 
        [0.5, 0.5, 0.5,],
        [0.0, 1.0, 0.0,],
        [0.5, 0.0, 0.0,],
        [0.0, 0.0, 0.5,],
        [0.5, 0.5, 0.0,],
        [0.5, 0.0, 0.5,],
        [1.0, 0.0, 0.0,],
        [0.75, 0.75, 0.75,],
        [0.0, 0.5, 0.5,],
        [1.0, 1.0, 1.0,], 
        [1.0, 1.0, 0.0,]])

pallet8 = np.array([
        [0.0, 0.0, 0.0,],
        [0.0, 0.0, 1.0,],
        [0.0, 1.0, 0.0,],
        [0.0, 1.0, 1.0,],
        [1.0, 0.0, 0.0,],
        [1.0, 0.0, 1.0,],
        [1.0, 1.0, 0.0,],
        [1.0, 1.0, 1.0,],])

colorPalletes = [pallet8, pallet16]
pallet1Gs = np.array([
        [0.0], 
        [1.0],])

pallet2Gs = np.array([
                [i / 3.0] for i in range(4)])

pallet4Gs = np.array([
        [i / 15.0] for i in range(16)])

gsPalletes = [pallet1Gs,pallet2Gs,pallet4Gs]

"""img = cv2.imread(pathSmall/filesSmall[1], cv2.IMREAD_COLOR_RGB)
img = imgToFloat(img)
floyd_dithered = FloydDithering(img, Pallet=pallet16)
plt.imshow(floyd_dithered)
plt.show()
quant = kwant_colorFit(img, Pallet=pallet16)
plt.imshow(quant)
plt.show()"""

document = Document()
document.add_heading('Raport 3',0) # tworzenie nagłówków, druga wartość to poziom nagłówka 
document.add_heading('Skalowanie małych zdjęć',1)

for i in tqdm(range(len(filesSmall)), desc="Processing small images"):
        document.add_heading('Zdjęcie {} '.format(filesSmall[i]), 2) 
        img = cv2.cvtColor(cv2.imread(pathSmall / filesSmall[i]),cv2.COLOR_BGR2RGB)
        img=imgToFloat(img)
        #add original photo to the document
        document.add_heading('Oryginał', 2)
        fig, axes = plt.subplots(1, 1, figsize=(10, 10))
        axes.imshow(img, cmap='gray')
        axes.set_title("Oryginał")
        axes.axis("off")
        plt.tight_layout()
        memfile = BytesIO()
        plt.savefig(memfile, format='png')
        memfile.seek(0)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        plt.close()
        for pallete in tqdm(range(2), desc="Processing color palletes"):
                document.add_heading(f'Paleta {(pallete+1)*8}', 2)#temp
                quantized = kwant_colorFit(img,colorPalletes[pallete])                
                rand_dithered = randDithering(img[:, :, 0])
                organized_dithered = organizedDithering(img, Pallet=colorPalletes[pallete])
                floyd_dithered = FloydDithering(img, Pallet=colorPalletes[pallete])
                
                # Plot the images
                fig, axes = plt.subplots(2, 2, figsize=(10, 10))
                axes = axes.ravel()
                axes[0].imshow(quantized)
                axes[0].set_title("Kwantyzacja")
                axes[0].axis("off")
                
                axes[1].imshow(rand_dithered,cmap='gray')
                axes[1].set_title("Random Dithering")
                axes[1].axis("off")
                
                axes[2].imshow(organized_dithered)
                axes[2].set_title("Zorganizowany Dithering")
                axes[2].axis("off")
                
                axes[3].imshow(floyd_dithered)
                axes[3].set_title("Floyd-Steinberg Dithering")
                axes[3].axis("off")
                
                plt.tight_layout()
                memfile = BytesIO()
                plt.savefig(memfile, format='png')
                memfile.seek(0)
                document.add_picture(memfile, width=Inches(6))
                memfile.close()
                plt.close(fig)


# Process grayscale images from filesGs
for i in tqdm(range(len(filesGs)), desc="Processing grayscale images"):
        document.add_heading('Zdjęcie {} '.format(filesGs[i]), 2)
        img = cv2.imread(str(pathGs / filesGs[i]), cv2.IMREAD_GRAYSCALE)
        img = imgToFloat(img)
        #add original photo to the document
        document.add_heading('Oryginał', 2)
        fig, axes = plt.subplots(1, 1, figsize=(10, 10))
        axes.imshow(img, cmap='gray')
        axes.set_title("Oryginał")
        axes.axis("off")
        plt.tight_layout()
        memfile = BytesIO()
        plt.savefig(memfile, format='png')
        memfile.seek(0)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        plt.close()
        for pallete in tqdm(range(len(gsPalletes)), desc="Processing grayscale palettes"):
                document.add_heading(f'Paleta {len(gsPalletes[pallete])} odcieni szarości', 2)
                quantized = kwant_colorFit(img, gsPalletes[pallete])
                rand_dithered = randDithering(img)
                organized_dithered = organizedDithering(img, Pallet=gsPalletes[pallete])
                floyd_dithered = FloydDithering(img, Pallet=gsPalletes[pallete])

                # Plot the images
                fig, axes = plt.subplots(2, 2, figsize=(10, 10))
                axes = axes.ravel()

                axes[0].imshow(quantized, cmap='gray')
                axes[0].set_title("Kwantyzacja")
                axes[0].axis("off")

                axes[1].imshow(rand_dithered, cmap='gray')
                axes[1].set_title("Random Dithering")
                axes[1].axis("off")

                axes[2].imshow(organized_dithered, cmap='gray')
                axes[2].set_title("Zorganizowany Dithering")
                axes[2].axis("off")

                axes[3].imshow(floyd_dithered, cmap='gray')
                axes[3].set_title("Floyd-Steinberg Dithering")
                axes[3].axis("off")

                plt.tight_layout()
                memfile = BytesIO()
                plt.savefig(memfile, format='png')
                memfile.seek(0)
                document.add_picture(memfile, width=Inches(6))
                memfile.close()
                plt.close(fig)

document.save('raport4.docx')


"""
#img = kwant_colorFit(img,pallet16)

#print(np.unique(FloydDithering(img,np.linspace(0,1,2).reshape(2,1))).size)"""

#if shape (1,1,3), pallete 16,3, else 1d array



