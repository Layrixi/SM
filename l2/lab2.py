import numpy as np
import matplotlib.pyplot as plt
import cv2
import pathlib
from docx import Document
from docx.shared import Inches
from io import BytesIO



def imgToUInt8(img):
    if np.issubdtype(img.dtype,np.float32):
        img = (img*255).astype(np.uint8)
    return img
def imgToFloat(img):
    if np.issubdtype(img.dtype,np.uint8):
        img = img.astype(np.float32)/255.
    return img

def subplot9(img,document=None):
    
    R,G,B = cv2.split(img)
    Y1=0.299 * R + 0.587 * G + 0.114 * B
    Y2=0.2126 * R + 0.7152 * G + 0.0722 * B

    plt.subplot(3, 3, 1)
    plt.imshow(img)
    plt.title('Oryginalny')

    plt.subplot(3, 3, 2)
    plt.imshow(Y1,cmap='gray')
    plt.title('Y1')

    plt.subplot(3, 3, 3)
    plt.imshow(Y2,cmap='gray')
    plt.title('Y2')


    #szare kolor
    plt.subplot(3, 3, 4)
    plt.imshow(R,cmap='gray')
    plt.title('R')

    plt.subplot(3, 3, 5)
    plt.imshow(G,cmap='gray')
    plt.title('G')

    plt.subplot(3, 3, 6)
    plt.imshow(B,cmap='gray')
    plt.title('B')

    #sam kolor
    Ronly = img.copy()
    Ronly[:,:,1] = 0
    Ronly[:,:,2] = 0
    Gonly = img.copy()
    Gonly[:,:,0] = 0
    Gonly[:,:,2] = 0
    Bonly = img.copy()
    Bonly[:,:,0] = 0
    Bonly[:,:,1] = 0

    plt.subplot(3, 3, 7)
    plt.imshow(Ronly)
    plt.title('R')

    plt.subplot(3, 3, 8)
    plt.imshow(Gonly)
    plt.title('G')

    plt.subplot(3, 3, 9)
    plt.imshow(Bonly)
    plt.title('B')
    if document is not None:
        memfile = BytesIO() # tworzenie bufora
        plt.savefig(memfile) # z zapis do bufora
        plt.close() # zamykanie wykresu
        document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku
        memfile.close()
    else:
        plt.show()
        
path = pathlib.Path().absolute() /'l2' / 'IMG_INTRO'
files=['A1.png','A2.jpg','A3.png','A4.jpg','B01.png','B02.jpg']
img1 = plt.imread(path/files[4])
img2 = cv2.imread(path/files[4])


print(img1.dtype)
print(img1.shape)
print(np.min(img1),np.max(img1))

print(img2.dtype)
print(img2.shape)
print(np.min(img2),np.max(img2))
#plt.imshow(img1 )
#plt.show()

#subplot9(img2)


#zad3
xlim = np.shape(img2)[1]
ylim = np.shape(img2)[0]
#fragments cutting

fragments = []  
for y in range(0, ylim, 200):
    for x in range(0, xlim, 200):
        fragment = img2[y:y+200, x:x+200].copy()
        fragments.append(fragment)

img3 = cv2.imread(path/files[5])
img3 = cv2.cvtColor(img3, cv2.COLOR_BGR2RGB)
xlim = np.shape(img3)[1]
ylim = np.shape(img3)[0]
for y in range(0, ylim, 200):
    for x in range(0, xlim, 200):
        fragment = img3[y:y+200, x:x+200].copy()
        fragments.append(fragment)


document = Document()
document.add_heading('Report 2 Jacek Jacewicz',0) # tworzenie nagłówków, druga wartość to poziom nagłówka 

for i,fragment in enumerate(fragments):
    document.add_heading('fragment nr - {} '.format(i),2)
    subplot9(fragment,document)
    

document.save('report2.docx') # zapis do pliku

