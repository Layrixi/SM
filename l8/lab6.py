import numpy as np
import matplotlib.pyplot as plt
import cv2
import os
import pathlib
from docx import Document
from docx.shared import Inches
from io import BytesIO
import sys
from tqdm import tqdm

class CompressedData:
    def __init__(self, compressed_data, original_shape):
        self.compressed_data = compressed_data
        self.original_shape = original_shape
        self.dataSize = (get_size(compressed_data) + get_size(original_shape))

    
def get_size(obj, seen=None):
    size = sys.getsizeof(obj)
    if seen is None:
        seen = set()
    obj_id = id(obj)
    if obj_id in seen:
        return 0
    # Important mark as seen *before* entering recursion to gracefully handle
    # self-referential objects
    seen.add(obj_id)
    if isinstance(obj,np.ndarray):
        size=obj.nbytes
    elif isinstance(obj, dict):
        size += sum([get_size(v, seen) for v in obj.values()])
        size += sum([get_size(k, seen) for k in obj.keys()])
    elif hasattr(obj, '__dict__'):
        size += get_size(obj.__dict__, seen)
    elif hasattr(obj, '__iter__') and not isinstance(obj, (str, bytes, bytearray)):
        size += sum([get_size(i, seen) for i in obj])
    return size

#zliczanie jak długo nie wystąpił ten sam symbol
def count_different_symbols(symbols, start):
    count =1
    for i in range(start+1, len(symbols)):
        if symbols[i] != symbols[i-1]:
            count += 1
        else:
            count-=1
            break
    return count
    

#zliczanie ile razy ten sam symbol
def count_repeated_symbols(symbols, start):
    count = 1
    while start + count < len(symbols) and symbols[start] == symbols[start + count]:
        count += 1
    return count

def rle_compress(data):
    original_shape = data.shape
    data = data.flatten()  # Flatten the 2D array into 1D
    compressed = np.zeros(len(data) * 2, dtype=int)
    index = 0
    i = 0

    while i < len(data):
        count = count_repeated_symbols(data, i)
        compressed[index] = count
        compressed[index + 1] = data[i]
        index += 2
        i += count

    return CompressedData(compressed[:index], original_shape)

def rle_decompress(compressed_data):
    decompressed = []
    compressed = compressed_data.compressed_data
    for i in range(0, len(compressed), 2):
        count = compressed[i]
        symbol = compressed[i + 1]
        decompressed.extend([symbol] * count)

    return np.array(decompressed, dtype=int).reshape(compressed_data.original_shape)

def ByteRun_compress(data):
    original_shape = data.shape
    data = data.flatten()  # Flatten the 2D array into 1D
    compressed = np.zeros(len(data) * 2, dtype=int)
    i = 0
    index = 0
    maxPackSize = 128

    while i < len(data):
        # If repetition
        if i + 1 < len(data) and data[i] == data[i + 1]:
            count = count_repeated_symbols(data, i)
            if count > maxPackSize:
                count = maxPackSize
            n = -count + 1
            compressed[index] = n
            compressed[index + 1] = data[i]
            index += 2
            i += count
        # If no repetition
        else:
            count = count_different_symbols(data, i)
            if count > maxPackSize:
                count = maxPackSize
            compressed[index] = count - 1
            for j in range(count):
                compressed[index + 1 + j] = data[i + j]
            index += count + 1
            i += count

    return CompressedData(compressed[:index],original_shape)

def ByteRun_decompress(compressed_data):
    decompressed = []
    i = 0
    compressed = compressed_data.compressed_data
    while i < len(compressed):
        n = compressed[i]
        if n < 0:
            # Repeated symbols
            count = -n + 1
            symbol = compressed[i + 1]
            decompressed.extend([symbol] * count)
            i += 2
        else:
            # Non-repeated symbols
            count = n + 1
            decompressed.extend(compressed[i + 1:i + 1 + count])
            i += 1 + count

    return np.array(decompressed, dtype=int).reshape(compressed_data.original_shape)


path = pathlib.Path().absolute() / 'l6'
files = ['document.png', 'kolor.png', 'techniczny.png']


#ByteRun test
data = np.array([
    [5, 5, 5, 1],
    [2, 3, 4, 4],
    [1, 2, 3, 4]
])
compressed_byterun = ByteRun_compress(data)
print("Original             : \n ",data)
print("Got                  : \n", compressed_byterun.compressed_data)
print("Expected             : \n [-2  5  2  1  2  3 -1  4  3  1  2  3  4]")
print("ByteRun Decompressed : \n",ByteRun_decompress(compressed_byterun))

#RLE test
data = np.array([1,1,1,1,1,2,2,2,3,4,5,6,6,6,6,1])
compressed_rle = rle_compress(data)
print("Original             : \n",data)
print("Got                  : \n", compressed_rle.compressed_data)
print("Expected             : \n [5 1 3 2 1 3 1 4 1 5 4 6 1 1]")
print("RLE Decompressed     : \n",rle_decompress(compressed_rle))
"""

document = Document()
document.add_heading('Raport 6',0) # tworzenie nagłówków, druga wartość to poziom nagłówka

document.add_heading('Dane testowe do kompresji', level=1)
cases = [
            np.array([1,1,1,1,2,1,1,1,1,2,1,1,1,1]),
            np.array([1,2,3,1,2,3,1,2,3]),
            np.array([5,1,5,1,5,5,1,1,5,5,1,1,5]),
            np.array([-1,-1,-1,-5,-5,-3,-4,-2,1,2,2,1]),
            np.zeros((1,520)),
            np.arange(0,521,1),
            np.eye(7),
            np.dstack([np.eye(7),np.eye(7),np.eye(7)]),
            np.ones((1,1,1,1,1,1,10))]
document.add_paragraph('W tabelach zamieszczono jedynie pierwsze 5 elementów, aby nie zajmować zbyt dużo miejsca.')
document.add_heading('Dane testowe do kompresji, algorytm RLE', level=2)
table = document.add_table(rows=len(cases)+1, cols=6)
table.style = 'Table Grid'
table.cell(0, 0).text = 'Dane oryginalne'
table.cell(0, 1).text = 'Dane skompresowane'
table.cell(0, 2).text = 'Dane zdekompresowane'
table.cell(0, 3).text = 'Rozmiar oryginalny'
table.cell(0, 4).text = 'Rozmiar skompresowany'
table.cell(0, 5).text = 'Czy oryginał = dekompresja'
for i, data in tqdm(enumerate(cases), desc='Processing', total=len(cases)):
    # Compression
    compressed = rle_compress(data)
    decompressed = rle_decompress(compressed)  # Pass original shape

    # Fill the table
    table.cell(i + 1, 0).text = str(data.flatten()[:5].tolist()) + ("..." if data.size > 5 else "")
    table.cell(i + 1, 1).text = str(compressed.compressed_data[:5].tolist()) + ("..." if len(compressed.compressed_data) > 5 else "")
    table.cell(i + 1, 2).text = str(decompressed.flatten()[:5].tolist()) + ("..." if decompressed.size > 5 else "")
    table.cell(i + 1, 3).text = str(get_size(data))
    table.cell(i + 1, 4).text = str(compressed.dataSize)
    table.cell(i + 1, 5).text = str(np.array_equal(data, decompressed))

document.add_heading('Dane testowe do kompresji, algorytm ByteRun', level=2)
table = document.add_table(rows=len(cases)+1, cols=6)
table.style = 'Table Grid'
table.cell(0, 0).text = 'Dane oryginalne'
table.cell(0, 1).text = 'Dane skompresowane'
table.cell(0, 2).text = 'Dane zdekompresowane'
table.cell(0, 3).text = 'Rozmiar oryginalny'
table.cell(0, 4).text = 'Rozmiar skompresowany'
table.cell(0, 5).text = 'Czy oryginał = dekompresja'
for i, data in tqdm(enumerate(cases), desc='Processing', total=len(cases)):
    # Compression
    compressed = ByteRun_compress(data)
    decompressed = ByteRun_decompress(compressed)  # Pass original shape

    # Fill the table
    table.cell(i + 1, 0).text = str(data.flatten()[:5].tolist()) + ("..." if data.size > 5 else "")
    table.cell(i + 1, 1).text = str(compressed.compressed_data[:5].tolist()) + ("..." if len(compressed.compressed_data) > 5 else "")
    table.cell(i + 1, 2).text = str(decompressed.flatten()[:5].tolist()) + ("..." if decompressed.size > 5 else "")
    table.cell(i + 1, 3).text = str(get_size(data))
    table.cell(i + 1, 4).text = str(compressed.dataSize)
    table.cell(i + 1, 5).text = str(np.array_equal(data, decompressed))


document.add_heading('Obrazy do kompresji', level=1)
images = []
for file in files:
    # Load the image
    img = cv2.imread(path / file)
    img = img.astype(int)  # Ensure the image is of type uint8
    images.append(img)
    memfile = BytesIO()  # tworzenie bufora
    is_success, buffer = cv2.imencode(".png", img)  # kodowanie obrazu do formatu PNG
    memfile.write(buffer)  # zapisanie zakodowanego obrazu do bufora
    memfile.seek(0)  # ustawienie wskaźnika na początek bufora
    document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku
    memfile.close()

table = document.add_table(rows=len(images)+1, cols=8)
table.style = 'Table Grid'
table.cell(0, 0).text = 'Nr zdjęcia'
table.cell(0, 1).text = 'Rozmiar oryginalny'
table.cell(0, 2).text = 'Rozmiar skompresowany RLE'
table.cell(0, 3).text = 'Rozmiar skompresowany ByteRun'
table.cell(0, 4).text = 'Stopień kompresji RLE'
table.cell(0, 5).text = 'Stopień kompresji ByteRun'
table.cell(0, 6).text = 'Czy oryginał = dekompresja RLE'
table.cell(0, 7).text = 'Czy oryginał = dekompresja ByteRun'

for i,image in tqdm(enumerate(images),desc ='Processing', total=len(images)):
    # Compression
    compressedRLE = rle_compress(image)
    decompressedRLE = rle_decompress(compressedRLE)  # Pass original shape
    compressedByteRun = ByteRun_compress(image)
    decompressedByteRun = ByteRun_decompress(compressedByteRun)  # Pass original shape

    table.cell(i + 1, 0).text = str('Obraz ' + str(i+1))
    table.cell(i + 1, 1).text = str(get_size(image))
    table.cell(i + 1, 2).text = str(compressedRLE.dataSize)
    table.cell(i + 1, 3).text = str(compressedByteRun.dataSize)
    table.cell(i + 1, 4).text = str(get_size(image)/compressedRLE.dataSize)
    table.cell(i + 1, 5).text = str(get_size(image)/compressedByteRun.dataSize)
    table.cell(i + 1, 6).text = str(np.array_equal(image, decompressedRLE))
    table.cell(i + 1, 7).text = str(np.array_equal(image, decompressedByteRun))


    
document.save(path / 'Raport6.docx')

#"""