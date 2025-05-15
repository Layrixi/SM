import numpy as np
import matplotlib.pyplot as plt
import cv2
import os
import pathlib
from docx import Document
from io import BytesIO
import sys
from tqdm import tqdm
import scipy.fftpack
from docx.shared import Inches


class ver2:
    def __init__(self, Y, Cb, Cr, OGShape, Ratio="4:4:4", QY=np.ones((8, 8)), QC=np.ones((8, 8))):
        self.Y = Y
        self.Cb = Cb
        self.Cr = Cr
        self.shape = OGShape  # Store the original shape of the image
        self.ChromaRatio = Ratio
        self.QY = QY
        self.QC = QC

def testShow(PRZED_RGB,PO_RGB,document=None):
    # Wyświetlenie obrazów przed i po kompresji
    fig, axs = plt.subplots(4, 2 , sharey=True   )
    fig.set_size_inches(9,13)
    # obraz oryginalny 
    axs[0,0].imshow(PRZED_RGB) #RGB 
    PRZED_YCrCb=cv2.cvtColor(PRZED_RGB,cv2.COLOR_RGB2YCrCb)
    axs[1,0].imshow(PRZED_YCrCb[:,:,0],cmap=plt.cm.gray) 
    axs[2,0].imshow(PRZED_YCrCb[:,:,1],cmap=plt.cm.gray)
    axs[3,0].imshow(PRZED_YCrCb[:,:,2],cmap=plt.cm.gray)

    # obraz po dekompresji
    axs[0,1].imshow(PO_RGB) #RGB 
    PO_YCrCb=cv2.cvtColor(PO_RGB,cv2.COLOR_RGB2YCrCb)
    axs[1,1].imshow(PO_YCrCb[:,:,0],cmap=plt.cm.gray)
    axs[2,1].imshow(PO_YCrCb[:,:,1],cmap=plt.cm.gray)
    axs[3,1].imshow(PO_YCrCb[:,:,2],cmap=plt.cm.gray)
    if document is not None:
        # Save the figure to a BytesIO object
        img_stream = BytesIO()
        plt.savefig(img_stream, format='png')
        img_stream.seek(0)
        # Add the image to the document
        document.add_picture(img_stream, width=Inches(6))
        # Close the BytesIO stream
        img_stream.close()
    else:
        plt.show()
    plt.close(fig)  # Close the figure to free memory

class CompressedData:
    def __init__(self, compressed_data, original_shape):
        self.compressed_data = compressed_data
        self.original_shape = original_shape
        self.dataSize = (get_size(compressed_data) + get_size(original_shape))

class CompressedJPEG:
    def __init__(self, S, OGShape):
        self.S = S
        self.OGShape = OGShape

def get_size(obj, seen=None):
    size = sys.getsizeof(obj)
    if seen is None:
        seen = set()
    obj_id = id(obj)
    if obj_id in seen:
        return 0
    # Important mark as seen *before* entering recursion to gracefully handle
    # self-referential objects
    seen.add(obj_id)
    if isinstance(obj,np.ndarray):
        size=obj.nbytes
    elif isinstance(obj, dict):
        size += sum([get_size(v, seen) for v in obj.values()])
        size += sum([get_size(k, seen) for k in obj.keys()])
    elif hasattr(obj, '__dict__'):
        size += get_size(obj.__dict__, seen)
    elif hasattr(obj, '__iter__') and not isinstance(obj, (str, bytes, bytearray)):
        size += sum([get_size(i, seen) for i in obj])
    return size
#zliczanie jak długo nie wystąpił ten sam symbol
def count_different_symbols(symbols, start):
    count =1
    for i in range(start+1, len(symbols)):
        if symbols[i] != symbols[i-1]:
            count += 1
        else:
            count-=1
            break
    return count

#zliczanie ile razy ten sam symbol
def count_repeated_symbols(symbols, start):
    count = 1
    while start + count < len(symbols) and symbols[start] == symbols[start + count]:
        count += 1
    return count

def rle_compress(data):
    original_shape = data.shape
    data = data.flatten()  # Flatten the 2D array into 1D
    compressed = np.zeros(len(data) * 2, dtype=int)
    index = 0
    i = 0

    while i < len(data):
        count = count_repeated_symbols(data, i)
        compressed[index] = count
        compressed[index + 1] = data[i]
        index += 2
        i += count

    return CompressedData(compressed[:index], original_shape)

def rle_decompress(compressed_data):
    decompressed = []
    compressed = compressed_data.compressed_data
    for i in range(0, len(compressed), 2):
        count = compressed[i]
        symbol = compressed[i + 1]
        decompressed.extend([symbol] * count)

    return np.array(decompressed, dtype=int).reshape(compressed_data.original_shape)

def ByteRun_compress(data):
    original_shape = data.shape  # Store the original shape of the input data
    data = data.flatten()  # Flatten the 2D array into 1D
    compressed = np.zeros(len(data) * 2, dtype=int)
    i = 0
    index = 0
    maxPackSize = 128

    while i < len(data):
        # If repetition
        if i + 1 < len(data) and data[i] == data[i + 1]:
            count = count_repeated_symbols(data, i)
            if count > maxPackSize:
                count = maxPackSize
            n = -count + 1
            compressed[index] = n
            compressed[index + 1] = data[i]
            index += 2
            i += count
        # If no repetition
        else:
            count = count_different_symbols(data, i)
            if count > maxPackSize:
                count = maxPackSize
            compressed[index] = count - 1
            for j in range(count):
                compressed[index + 1 + j] = data[i + j]
            index += count + 1
            i += count

    return compressed[:index], original_shape  # Return both compressed data and original shape


def ByteRun_decompress(compressed_data, original_shape):
    decompressed = []
    i = 0
    compressed = compressed_data
    while i < len(compressed):
        n = compressed[i]
        if n < 0:
            # Repeated symbols
            count = -n + 1
            symbol = compressed[i + 1]
            decompressed.extend([symbol] * count)
            i += 2
        else:
            # Non-repeated symbols
            count = n + 1
            decompressed.extend(compressed[i + 1:i + 1 + count])
            i += 1 + count

    return np.array(decompressed, dtype=int).reshape(original_shape)

def zigzag(A):
    template= np.array([
            [0,  1,  5,  6,  14, 15, 27, 28],
            [2,  4,  7,  13, 16, 26, 29, 42],
            [3,  8,  12, 17, 25, 30, 41, 43],
            [9,  11, 18, 24, 31, 40, 44, 53],
            [10, 19, 23, 32, 39, 45, 52, 54],
            [20, 22, 33, 38, 46, 51, 55, 60],
            [21, 34, 37, 47, 50, 56, 59, 61],
            [35, 36, 48, 49, 57, 58, 62, 63],
            ])
    if len(A.shape)==1:
        B=np.zeros((8,8))
        for r in range(0,8):
            for c in range(0,8):
                B[r,c]=A[template[r,c]]
    else:
        B=np.zeros((64,))
        for r in range(0,8):
            for c in range(0,8):
                B[template[r,c]]=A[r,c]
    return B

#przed DCT trzeba wycentrować dane

def dct2(a):
    return scipy.fftpack.dct( scipy.fftpack.dct( a.astype(float), axis=0, norm='ortho' ), axis=1, norm='ortho' )

def idct2(a):
    return scipy.fftpack.idct( scipy.fftpack.idct( a.astype(float), axis=0 , norm='ortho'), axis=1 , norm='ortho')

def CompressBlock(block,Q):
    dct2block=dct2(block)
    #Kwantyzacja poprzez zastąpienie wartości zmiennoprzecinkowych, wartościami całkowitymi
    #przy pomocy macierzy kwantyzacji
    qd=np.round(dct2block/Q).astype(int)
    vector=zigzag(qd)
    return vector

def DecompressBlock(vector,Q):
    #dezigzag
    dezig=zigzag(vector)
    pd=dezig*Q
    block=idct2(pd)
    return block

## podział na bloki
# L - warstwa kompresowana
# S - wektor wyjściowy
def CompressLayer(L,Q):
    S = np.array([])
    for w in range(0, L.shape[0], 8):
        for k in range(0, L.shape[1], 8):
            block = L[w:(w + 8), k:(k + 8)]
            # Center the data
            block = block - 128
            S = np.append(S, CompressBlock(block, Q))
    # Lossless compression
    S, original_shape = ByteRun_compress(S)
    return S, original_shape


def DecompressLayer(S, original_shape, Q):
    # Lossless decompression
    S = ByteRun_decompress(S, original_shape)
    # Declare a matrix of the appropriate size
    height, width = 128, 128
    L = np.zeros((height, width))
    idx = 0
    for w in range(0, height, 8):
        for k in range(0, width, 8):
            vector = S[idx:(idx + 64)]
            idx += 64
            L[w:(w + 8), k:(k + 8)] = DecompressBlock(vector, Q)
    # Decenter the data
    L += 128
    return L


def CompressJPEG(RGB, Ratio="4:4:4", QY=np.ones((8, 8)), QC=np.ones((8, 8)), ratio="4:4:4"):
    # RGB -> YCrCb
    YCrCb = cv2.cvtColor(RGB, cv2.COLOR_RGB2YCrCb).astype(int)
    JPEG = ver2(YCrCb[:, :, 0], YCrCb[:, :, 1], YCrCb[:, :, 2], RGB.shape, Ratio, QY, QC)
    
    # Chroma subsampling
    if ratio == "4:2:2":
        JPEG.Cb = JPEG.Cb[::2, :]
        JPEG.Cr = JPEG.Cr[::2, :]
    elif ratio == "4:2:0":
        pass
    else:  # Default "4:4:4"
        JPEG.Cb = JPEG.Cb
        JPEG.Cr = JPEG.Cr
    
    # Compress each layer
    JPEG.Y, JPEG.Y_shape = CompressLayer(JPEG.Y, JPEG.QY)
    JPEG.Cr, JPEG.Cr_shape = CompressLayer(JPEG.Cr, JPEG.QC)
    JPEG.Cb, JPEG.Cb_shape = CompressLayer(JPEG.Cb, JPEG.QC)

    return JPEG


def DecompressJPEG(JPEG):
    # Decompress each layer
    Y = DecompressLayer(JPEG.Y, JPEG.Y_shape, JPEG.QY)
    Cr = DecompressLayer(JPEG.Cr, JPEG.Cr_shape, JPEG.QC)
    Cb = DecompressLayer(JPEG.Cb, JPEG.Cb_shape, JPEG.QC)

    # Chroma resampling
    if JPEG.ChromaRatio == "4:2:2":
        # Resample horizontally to match the original width
        Cb = np.repeat(Cb, 2, axis=1)[:JPEG.shape[0], :JPEG.shape[1]]
        Cr = np.repeat(Cr, 2, axis=1)[:JPEG.shape[0], :JPEG.shape[1]]
    elif JPEG.ChromaRatio == "4:2:0":
        pass
    else:  # Default "4:4:4"
        # No resampling needed
        Cb = Cb[:JPEG.shape[0], :JPEG.shape[1]]
        Cr = Cr[:JPEG.shape[0], :JPEG.shape[1]]

    # Reconstruct the image
    YCrCb = np.zeros((JPEG.shape[0], JPEG.shape[1], 3))
    YCrCb[:, :, 0] = Y
    YCrCb[:, :, 1] = Cb
    YCrCb[:, :, 2] = Cr

    # YCrCb -> RGB
    return cv2.cvtColor(YCrCb.astype(np.uint8), cv2.COLOR_YCrCb2RGB)

#irreversible part


#ByteRun test
data = np.array([
    [5, 5, 5, 1],
    [2, 3, 4, 4],
    [1, 2, 3, 4]
])
QY= np.array([
        [16, 11, 10, 16, 24,  40,  51,  61],
        [12, 12, 14, 19, 26,  58,  60,  55],
        [14, 13, 16, 24, 40,  57,  69,  56],
        [14, 17, 22, 29, 51,  87,  80,  62],
        [18, 22, 37, 56, 68,  109, 103, 77],
        [24, 36, 55, 64, 81,  104, 113, 92],
        [49, 64, 78, 87, 103, 121, 120, 101],
        [72, 92, 95, 98, 112, 100, 103, 99],
        ])
QC= np.array([
        [17, 18, 24, 47, 99, 99, 99, 99],
        [18, 21, 26, 66, 99, 99, 99, 99],
        [24, 26, 56, 99, 99, 99, 99, 99],
        [47, 66, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        ])

QN= np.ones((8,8))

#przed wywołaniem jpega trzeba konwert z bgr na rgb
path = pathlib.Path().absolute() / 'l8'
filesFull = ['burnice_Full.png', 'cantarella_Full.png', 'cantarella_Full2.png', 'cantarella_Full3.png']
files = ['burnice1.png','burnice2.png','burnice3.png','cantarella1.png',
         'cantarella2.png','cantarella3.png','cantarella4.png',
         'cantarella5.png','cantarella6.png','cantarella7.png',
         'cantarella8.png','cantarella9.png',
         'cantarella10.png','cantarella11.png','cantarella12.png']
document = Document()
document.add_heading('Raport 8 Jacek Jacewicz', level=1)
for i, file in enumerate(files):
    if i == 0:
        # Add burnice full
        img_path = path / filesFull[0]
        document.add_heading('Zdjęcie pełne nr 1', level=2)
        document.add_picture(str(img_path), width=Inches(6))  # Set width to 6 inches
    elif i == 3:
        # Add cantarellaFull
        img_path = path / filesFull[1]
        document.add_heading('Zdjęcie pełne nr 2', level=2)
        document.add_picture(str(img_path), width=Inches(6))  # Set width to 6 inches
    elif i == 7:
        # Add cantarellaFull2
        img_path = path / filesFull[2]
        document.add_heading('Zdjęcie pełne nr 3', level=2)
        document.add_picture(str(img_path), width=Inches(6))  # Set width to 6 inches
    elif i == 11:
        # Add cantarellaFull3
        img_path = path / filesFull[3]
        document.add_heading('Zdjęcie pełne nr 4', level=2)
        document.add_picture(str(img_path), width=Inches(6))  # Set width to 6 inches
    document.add_heading('fragmenty przed i po kompresji JPEG:', level=2)
    #do sprawdzenia: 4:4:4, 4:2:2, tablice kwantyzacji albo tablice 1
    img_path = path / file
    img = cv2.imread(img_path)
    img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
    #4:4:4 i tablice kwantyzacji
    JPEG = CompressJPEG(img, Ratio="4:4:4", QY=QY, QC=QC)
    document.add_heading('4:4:4 i tablice kwantyzacji', level=3)
    PO_RGB = DecompressJPEG(JPEG)
    testShow(img, PO_RGB, document)    
    #4:4:4 i tablice 1
    JPEG2 = CompressJPEG(img, Ratio="4:4:4", QY=QN, QC=QN)
    document.add_heading('4:4:4 i tablice 1', level=3)
    PO_RGB = DecompressJPEG(JPEG2)
    testShow(img, PO_RGB, document)
    #4:2:2 i tablice 1
    JPEG3 = CompressJPEG(img, Ratio="4:2:2", QY=QN, QC=QN)
    document.add_heading('4:2:2 i tablice 1', level=3)
    PO_RGB = DecompressJPEG(JPEG3)
    testShow(img, PO_RGB, document)
    #4:2:2 i tablice kwantyzacji
    JPEG4 = CompressJPEG(img, Ratio="4:2:2", QY=QY, QC=QC)
    document.add_heading('4:2:2 i tablice kwantyzacji', level=3)
    PO_RGB = DecompressJPEG(JPEG4)
    testShow(img, PO_RGB, document)
    

document.save('raport.docx')
