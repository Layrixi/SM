import numpy as np
import matplotlib.pyplot as plt
import cv2
import pathlib
from docx import Document
from docx.shared import Inches
from io import BytesIO
from os import path
from tqdm import tqdm



def displayImParts(img, part_size=0.25, num_parts=3, document=None):
    height, width = img.shape[:2]
    part_height = int(height * part_size)
    part_width = int(width * part_size)
    
    fig, axes = plt.subplots(1, num_parts, figsize=(num_parts * 4, 4))
    if num_parts == 1:
        axes = [axes] 

    for i in range(num_parts):
        start_row = (i // (1 / part_size)) * part_height
        start_col = (i % (1 / part_size)) * part_width
        end_row = start_row + part_height
        end_col = start_col + part_width
        start_col = int(start_col)
        start_row = int(start_row)
        end_col = int(end_col)
        end_row = int(end_row)

        part_img = img[start_row:end_row, start_col:end_col]
        axes[i].imshow(part_img)
        axes[i].axis('off')
        axes[i].set_title(f'Wycinek {i+1}')

    plt.tight_layout()

    if document is not None:
        memfile = BytesIO()
        plt.savefig(memfile, format='png')
        memfile.seek(0)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
    else:
        plt.show()

    plt.close(fig)

def scaleInterpNearest(img,k):
    height, width = img.shape[:2]
    new_height = int(height * k)
    new_width = int(width * k)
    x_coords = np.linspace(0, height - 1, new_height)
    y_coords = np.linspace(0, width - 1, new_width)
    
    scaled_img = np.zeros((new_height, new_width, img.shape[2]), dtype=img.dtype)
    
    for i, x in enumerate(x_coords):
        for j, y in enumerate(y_coords):
            x0=np.floor(x)
            x1=np.ceil(x)
            y0=np.floor(y)
            y1=np.ceil(y)
            deltax = x - x0
            deltay = y - y0
            scaled_img[i, j] = (
                img[int(x0), int(y0)] * (1 - deltax) * (1 - deltay) +
                img[int(x1), int(y0)] * deltax * (1 - deltay) +
                img[int(x0), int(y1)] * (1 - deltax) * deltay +
                img[int(x1), int(y1)] * deltax * deltay
            )

            #scaled_img[i, j] = img[round(x), round(y)]
    
    return scaled_img

#k - mnożnik. 200%: k=2 itp.
def scaleNearest(img,k):
    height, width = img.shape[:2]
    new_height = int(height * k)
    new_width = int(width * k)
    x_coords = np.linspace(0, height - 1, new_height)
    y_coords = np.linspace(0, width - 1, new_width)
    
    scaled_img = np.zeros((new_height, new_width, img.shape[2]), dtype=img.dtype)
    
    for i, x in enumerate(x_coords):
        for j, y in enumerate(y_coords):
            scaled_img[i, j] = img[round(x), round(y)]
    
    return scaled_img

def lowresMean(img,k):
    step=int(np.ceil(1/np.sqrt(k)))

    height, width = img.shape[:2]
    new_height = int(height / step)
    new_width = int(width / step)

    scaled_img = np.zeros((new_height, new_width, img.shape[2]), dtype=img.dtype)
    
    for i in range(new_height):
        for j in range(new_width):
            block = img[i*step:(i+1)*step, j*step:(j+1)*step]
            scaled_img[i, j] = np.mean(block, axis=(0, 1))
    
    return scaled_img

def lowresWMean(img, k):
    step = int(np.ceil(1 / np.sqrt(k)))
    height, width = img.shape[:2]
    new_height = int(height / step)
    new_width = int(width / step)

    scaled_img = np.zeros((new_height, new_width, img.shape[2]), dtype=img.dtype)

    for i in range(new_height):
        for j in range(new_width):
            block = img[i * step:(i + 1) * step, j * step:(j + 1) * step]
            weights = np.random.rand(*block.shape)

            # Compute weighted mean for the block using vectorized operations
            weighted_sum = np.sum(block * weights, axis=(0, 1))
            sum_weights = np.sum(weights, axis=(0, 1))
            weighted_mean = weighted_sum / sum_weights

            scaled_img[i, j] = weighted_mean

    return scaled_img

#k - mnożnik. Jeśli zmniejszyć obraz o połowe to k = 0.5, jeśli do 10% to k = 0.1 itpp.
def lowresMedian(img,k):
    step=int(np.ceil(1/np.sqrt(k)))
    height, width = img.shape[:2]
    new_height = int(height / step)
    new_width = int(width / step)

    scaled_img = np.zeros((new_height, new_width, img.shape[2]), dtype=img.dtype)
    
    for i in range(new_height):
        for j in range(new_width):
            block = img[i*
                        step:(i+1)*step, j*step:(j+1)*step]
            scaled_img[i, j] = np.median(block, axis=(0, 1))
    
    return scaled_img

pathSmall = pathlib.Path().absolute() /'l3' / 'IMG_SMALL'
pathBig = pathlib.Path().absolute() /'l3' / 'IMG_BIG'

filesSmall=['SMALL_0001.tif',
       'SMALL_0002.png',
       'SMALL_0003.png',
        'SMALL_0010.jpg']
filesBig=['BIG_0001.jpg',
          'BIG_0002.jpg'
          ]

imgtest= np.zeros((3,3,3),dtype=np.uint8)
imgtest[1,1,:]=255
#img = cv2.imread(pathBig/filesBig[1])



document = Document()
document.add_heading('Raport 3',0) # tworzenie nagłówków, druga wartość to poziom nagłówka 
document.add_heading('Skalowanie małych zdjęć',1)


#scale up 4 small images
for i in tqdm(range(4), desc="Processing small images"):
    document.add_heading('Zdjęcie {} '.format(filesSmall[i]), 2) 
    img = cv2.cvtColor(cv2.imread(pathSmall / filesSmall[i]),cv2.COLOR_BGR2RGB)
    document.add_heading('Oryginał', 2)
    displayImParts(img, document=document)
    for k in tqdm([1, 0.5, 3], desc=f"Scaling small image {filesSmall[i]}"):
        document.add_heading(f'Metoda najbliższych sąsiadów, k={k}, zdjęcie {filesSmall[i]}', 2)
        scaled = scaleNearest(img, k)
        displayImParts(scaled, document=document)
        document.add_heading(f'Metoda interpolacji dwuliniowej, k={k}, zdjęcie {filesSmall[i]}', 2) 
        scaled = scaleInterpNearest(img, k)
        displayImParts(scaled, document=document)

document.add_heading('Skalowanie dużych zdjęć', 1)
#scale down 2 big images
for i in tqdm(range(2), desc="Processing big images"):
    document.add_heading('Zdjęcie {} '.format(filesBig[i]), 2) 
    img = cv2.cvtColor(cv2.imread(pathBig / filesBig[i]),cv2.COLOR_BGR2RGB)
    document.add_heading('Oryginał', 2)
    displayImParts(img, document=document)
    for k in tqdm([1.5, 0.07, 0.05], desc=f"Scaling big image {filesBig[i]}"):
        document.add_heading(f'Metoda średniej, k={k}, zdjęcie {filesBig[i]}', 2)
        scaled = lowresMean(img, k)
        displayImParts(scaled, document=document)
        document.add_heading(f'Metoda średniej ważonej, k={k}, zdjęcie {filesBig[i]}', 2) 
        scaled = lowresWMean(img, k)
        displayImParts(scaled, document=document)
        document.add_heading(f'Metoda mediany, k={k}, zdjęcie {filesBig[i]}', 2) 
        scaled = lowresMedian(img, k)
        displayImParts(scaled, document=document)

document.save('report3.docx') # zapis do pliku"""
