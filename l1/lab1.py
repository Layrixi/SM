import numpy as np
import matplotlib.pyplot as plt
import scipy.fftpack
import sounddevice as sd
import soundfile as sf
import pathlib
from docx import Document
from docx.shared import Inches
from io import BytesIO

#t=n/fs
def plotAudio(signal,fs,fsize,document=None,TimeMargin=[0,0.02]):
    fig ,axs = plt.subplots(2,1,figsize=(10,7)) # tworzenie plota

    plt.subplot(2,1,1)
    plt.xlim(TimeMargin)
    plt.xlabel("T(s)")
    plt.ylabel("Amplitude")
    plt.plot(np.arange(0,signal.shape[0])/fs,signal)
    plt.subplot(2,1,2)
    yf = scipy.fftpack.fft(signal,fsize)
    xf=np.arange(0,fs/2,fs/fsize)
    dbspectrum=20*np.log10( np.abs(yf[:fsize//2]))
    plt.xlabel("Fs(Hz)")
    plt.ylabel("Magnitude(dB)")
    plt.plot(np.arange(0,fs/2,fs/fsize),dbspectrum)
    max_magnitude = np.max(dbspectrum)
    max_frequency = xf[np.argmax(dbspectrum)]
    if document is not None:
        fig.suptitle('Time margin {}'.format(TimeMargin)) # Tytuł wykresu
        fig.tight_layout(pad=1.5) # poprawa czytelności 
        memfile = BytesIO() # tworzenie bufora
        fig.savefig(memfile) # z zapis do bufora
        document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku
        document.add_paragraph('Max db: {} dB w fs: {} Hz'.format(max_magnitude,max_frequency))
        memfile.close()



if __name__ == "__main__":
    path = pathlib.Path().absolute() / 'l1'/ 'SIN'
    z1path = pathlib.Path().absolute() / 'l1' / 'SOUND_INTRO' / 'sound1.wav'
    data,fs = sf.read(z1path, dtype=np.int32)
    #z1
    """left=np.zeros_like(data)
    right=np.zeros_like(data)
    left[:,0]=data[:,0]
    right[:,1]=data[:,1]
    mono = (left[:,0] + right[:,1]) / 2
    sf.write('lewy.wav', left, fs)
    sf.write('prawy.wav', right, fs)
    sf.write('mono.wav', mono, fs)"""


    files=['sin_60Hz.wav','sin_440Hz.wav','sin_8000Hz.wav']
    fsize=[2**8,2**12,2**16]
    document = Document()
    document.add_heading('Zmień ten tytuł',0) # tworzenie nagłówków, druga wartość to poziom nagłówka 
    Margins=[[0,0.02],[0.133,0.155]]
    #file = path / files[0]
    #data,fs = sf.read(file, dtype=np.int32)
    #plotAudio(data,fs,fsize[0],document,Margins[0])
    for file in files:
        document.add_heading('Plik - {}'.format(file),2)
        for i,Margin in enumerate(fsize):
            document.add_heading('Time margin {}'.format(Margin),3) # nagłówek sekcji, mozę być poziom wyżej
        
            ############################################################
            # Tu wykonujesz jakieś funkcje i rysujesz wykresy
            ############################################################
            data, fs = sf.read(path/file, dtype=np.int32)

            plotAudio(data,fs,Margin,document)



            
            
        
            
            ############################################################
            # Tu dodajesz dane tekstowe - wartosci, wyjscie funkcji ect.

            ############################################################

    document.save('report.docx') # zapis do pliku"""    

        