import numpy as np
import matplotlib.pyplot as plt
import scipy.fftpack
import sounddevice as sd
import soundfile as sf
import pathlib
from docx import Document
from docx.shared import Inches
from io import BytesIO
from scipy.interpolate import interp1d

#t=n/fs
def plotAudio(signal,fs,fsize,document=None,TimeMargin=[0,0.02],title="Audio"):
        fig ,axs = plt.subplots(2,1,figsize=(10,7)) # tworzenie plota
        fig.suptitle(title) 
        plt.subplot(2,1,1)
        plt.xlim(TimeMargin)
        plt.ylim(np.min(signal), np.max(signal)) 
        plt.xlabel("T(s)")
        plt.ylabel("Amplitude")
        plt.plot(np.arange(0,signal.shape[0])/fs,signal)


        plt.subplot(2,1,2)
        yf = scipy.fftpack.fft(signal,fsize)
        xf=np.arange(0,fs/2,fs/fsize)
        dbspectrum = 20 * np.log10(np.abs(yf[:fsize // 2])+np.finfo(np.float64).eps)
        plt.xlabel("Fs(Hz)")
        plt.ylabel("Magnitude(dB)")
        plt.plot(np.arange(0,fs/2,fs/fsize),dbspectrum)
        max_magnitude = np.max(dbspectrum)
        max_frequency = xf[np.argmax(dbspectrum)]
        if document is not None:
                fig.tight_layout(pad=1.5) # poprawa czytelności 
                memfile = BytesIO() # tworzenie bufora
                fig.savefig(memfile) # z zapis do bufora
                document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku
                document.add_paragraph('Max db: {} dB w fs: {} Hz'.format(max_magnitude,max_frequency))
                memfile.close()
        else:
                plt.show()
        plt.close()
def kwant(data, bit):
        d = 2**bit - 1
        
        if np.issubdtype(data.dtype, np.floating):
                m = np.min(data)
                n = np.max(data)
        else:
                m = np.iinfo(data.dtype).min
                n = np.iinfo(data.dtype).max

        #1
        DataF = data.astype(float)
        DataF = (DataF - m) / (n - m)

        #2
        DataF = np.round(DataF * d)
        DataF = DataF / d

        #3
        DataF = DataF * (n - m) + m

        return DataF.astype(data.dtype)
def decimSampReduction(data, fs, factor):
        return data[::factor],fs//factor
def interpSampReduction(data, fs,newFs, method='linear'):
        N=data.__len__()
        duration = N/fs
        N1=int(duration*newFs)
        t=np.linspace(0, duration, N)
        timeMoments = np.linspace(0, duration, N1)
        data=np.array(data)
        if method == 'linear':
                metode=interp1d(t,data,kind='linear')
        else:
                metode=interp1d(t,data, kind='cubic')

        data_new=metode(timeMoments).astype(data.dtype)
        return data_new, newFs

path = pathlib.Path().absolute() / 'l5'/ 'SIN'
singPath = pathlib.Path().absolute() / 'l5' / 'SING' 

files=['sin_60Hz.wav','sin_440Hz.wav','sin_8000Hz.wav','sin_combined.wav']
singFiles=['sing_high1.wav','sing_low1.wav','sing_medium1.wav']
fsize=2**16

"""file = pathlib.Path().absolute() / 'l5' / 'wicked.mp3'
data,fs = sf.read(file, dtype=np.int32)
wickedRedu,wickedFs = decimSampReduction(data, fs, 40)
sf.write(pathlib.Path().absolute() / 'l5' / f'wickedRedu.wav', wickedRedu, wickedFs)

#"""

document = Document()
document.add_heading('Raport 5',0) # tworzenie nagłówków, druga wartość to poziom nagłówka
#plotting
quantBits= [4,8,16,24]
decimParameters = [2,4,6,10,24]
interpParameters = [2000,4000,8000,11999,16000,16953,24000,41000]
sinTimeMargins = [[0,0.06],[0,0.04],[0,0.002],[0,0.01]]
for file in files:
        data, fs = sf.read(path/file, dtype=np.int32)
        timeMargin = sinTimeMargins[files.index(file)]
        document.add_heading(f'Plik:{file}, fs = {fs}', level=1) # dodanie tytułu do pliku
        plotAudio(data, fs, fsize, TimeMargin=timeMargin,title = "Original Audio",document=document)
        #plotAudio(data, fs, fsize, TimeMargin=timeMargin,title = "Original Audio")

        for i in range(len(quantBits)):
                dataQuant = kwant(data, quantBits[i])
                #plotAudio(dataQuant, fs, fsize, TimeMargin=timeMargin, title = (f'Quantized Audio, {quantBits[i]} bits'))
                plotAudio(dataQuant, fs, fsize, TimeMargin=timeMargin, title = (f'Quantized Audio, {quantBits[i]} bits'), document=document)

        for i in range(len(decimParameters)):
                dataDecim, fsDecim = decimSampReduction(data, fs, decimParameters[i])
                #plotAudio(dataDecim, fsDecim, fsize, TimeMargin=timeMargin, title = (f'Decimated Audio, {decimParameters[i]}'))
                plotAudio(dataDecim, fsDecim, fsize, TimeMargin=timeMargin, title = (f'Decimated Audio, parameter = {decimParameters[i]}, Fs after resampling: {fsDecim}Hz, Fs before resampling {fs}'), document=document)

        for i in range(len(interpParameters)):
                dataInterp, fsInterp = interpSampReduction(data, fs, interpParameters[i], method='linear')
                #plotAudio(dataInterp, fsInterp, fsize, TimeMargin=timeMargin, title = (f'Interpolated Audio, {interpParameters[i]} Hz, method=linear'))
                plotAudio(dataInterp, fsInterp, fsize, TimeMargin=timeMargin, title = (f'Interpolated Audio, {interpParameters[i]} Hz, Fs after resampling: {fsInterp}Hz, method=linear, Fs before resampling: {fs}'), document=document)
                dataInterp, fsInterp = interpSampReduction(data, fs, interpParameters[i], method='cubic')
                #plotAudio(dataInterp, fsInterp, fsize, TimeMargin=timeMargin, title = (f'Interpolated Audio, {interpParameters[i]} Hz, method=cubic'))
                plotAudio(dataInterp, fsInterp, fsize, TimeMargin=timeMargin, title = (f'Interpolated Audio, {interpParameters[i]} Hz, Fs after resampling: {fsInterp}Hz, method=cubic, Fs before resampling: {fs}'), document=document)
#"""

"""
#sing files
quantBits= [4,8]
decimParameters = [4,6,10,24]
interpParameters = [4000,8000,11999,16000,16953]
saveMode = True
for file in singFiles:
        data, fs = sf.read(singPath/file, dtype=np.int32)
        for i in range(len(quantBits)):
                dataQuant = kwant(data, quantBits[i])
                if saveMode:
                        sf.write(pathlib.Path().absolute() / 'l5' / 'quant' / f"{file}_quant{quantBits[i]}.wav", dataQuant, fs)
                
        for i in range(len(decimParameters)):
                dataDecim, fsDecim = decimSampReduction(data, fs, decimParameters[i])
                if saveMode:
                        sf.write(pathlib.Path().absolute() / 'l5' / 'decim' / f"{file}_decim{decimParameters[i]}.wav", dataDecim, fsDecim)
        for i in range(len(interpParameters)):
                dataInterp, fsInterp = interpSampReduction(data, fs, interpParameters[i], method='linear')
                if saveMode:
                        sf.write(pathlib.Path().absolute() / 'l5' / 'interp' / f"{file}_interp{interpParameters[i]}_linear.wav", dataInterp, fsInterp)
                dataInterp, fsInterp = interpSampReduction(data, fs, interpParameters[i], method='cubic')
                if saveMode:
                        sf.write(pathlib.Path().absolute() / 'l5' / 'interp' / f"{file}_interp{interpParameters[i]}_cubic.wav", dataInterp, fsInterp)"""
document.save(pathlib.Path().absolute() / 'l5' / 'Raport5.docx')
"""
#test 
data = np.arange(np.iinfo(np.int32).min, np.iinfo(np.int32).max, 1000, dtype=np.int32)
# Kwantyzacja na 5 bitów
kwant_data = Kwant(data, 4)
# Wyświetlenie wykresu
plt.plot(data, label='Orig')
plt.plot(kwant_data, label='Kwant')
plt.legend()
plt.title('Kwantyzacja 4 bitowa test')
plt.xlabel('Index')
plt.ylabel('Wartość')
plt.show()#"""



