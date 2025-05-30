from io import BytesIO
import pathlib
import random
import sys
import cv2
import numpy as np
import matplotlib.pyplot as plt
from tqdm import tqdm
from docx import Document
from docx.shared import Inches
from skimage.metrics import structural_similarity as ssim
import math
from scipy.stats import pearsonr, spearmanr
from scipy.optimize import curve_fit

#dopasowywanie korelacji
def inverse_func(x, a, b):
    return a / x + b

def linear_func(x, a, b):
    return a * x + b

def log_func(x, a, b):
    return a * np.log(x) + b

def fit_and_plot(param_values, metric_values, metric_name="PSNR", method_name="blur"):
    x = np.array(param_values)
    y = np.array(metric_values)
    
    # Oblicz korelacje
    pearson_corr, _ = pearsonr(x, y)
    spearman_corr, _ = spearmanr(x, y)

    # Dopasuj różne modele
    popt_inv, _ = curve_fit(inverse_func, x, y)
    popt_lin, _ = curve_fit(linear_func, x, y)
    popt_log, _ = curve_fit(log_func, x, y)

    # Generuj dane do wykresów
    x_fit = np.linspace(min(x), max(x), 100)
    y_inv = inverse_func(x_fit, *popt_inv)
    y_lin = linear_func(x_fit, *popt_lin)
    y_log = log_func(x_fit, *popt_log)

    # Rysuj
    plt.figure(figsize=(10, 6))
    plt.plot(x, y, 'o', label='Dane')
    plt.plot(x_fit, y_inv, '-', label='Dopasowanie odwrotne')
    plt.plot(x_fit, y_lin, '--', label='Dopasowanie liniowe')
    plt.plot(x_fit, y_log, '-.', label='Dopasowanie logarytmiczne')
    plt.xlabel("Parametr degradacji")
    plt.ylabel(metric_name)
    plt.title(f"{metric_name} vs parametr degradacji ({method_name})")
    plt.legend()
    plt.grid(True)
    plt.show()

    return {
        "pearson": pearson_corr,
        "spearman": spearman_corr,
        "inv_params": popt_inv,
        "lin_params": popt_lin,
        "log_params": popt_log
    }


def mse(K, I):
    return np.mean((I - K) ** 2)

def nmse(K, I):
    numerator = np.sum((I - K) ** 2)
    denominator = np.sum(K ** 2)
    return numerator / denominator

def psnr(K, I, max_I=255.0):
    mse_val = mse(K, I)
    #if no difference
    if mse_val == 0:
        return float('inf')
    return 10 * math.log10((max_I ** 2) / mse_val)

def image_fidelity(K, I):
    numerator = np.sum((K - I) ** 2)
    denominator = np.sum(K  * I)
    return 1 - (numerator / denominator)

def compute_ssim(K, I):
    return ssim(K, I, data_range=I.max() - I.min(), channel_axis=-1, win_size=3)

def szum(img,alpha,Noise_range=[-25,25]):
    rand = (Noise_range[1]-Noise_range[0])*np.random.random((img.shape))+Noise_range[0]
    noisy3 = (img + alpha * rand).clip(0,255).astype(np.uint8)
    return noisy3

def JPEGCompress(img, quality = 95):
    encode_param = [int(cv2.IMWRITE_JPEG_QUALITY), quality]
    result, encimg = cv2.imencode('.jpg', img, encode_param)
    decimg = cv2.imdecode(encimg, 1)
    return decimg

def plotSM(img1, img2, title1='Image 1', title2='Image 2',document=None):
    fig, axs = plt.subplots(1, 2 , sharey=True   )
    axs[0].set_title(title1)
    axs[1].set_title(title2)
    axs[0].imshow(img1)
    axs[1].imshow(img2)
    if document is not None:
        buf = BytesIO()
        plt.savefig(buf, format='png')
        buf.seek(0)
        document.add_picture(buf, width=Inches(5.0))
        buf.close()
    else:
        plt.show()
    plt.close(fig)

def add_results_table(document, title, results):
    document.add_heading(title, level=2)
    table = document.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parametry'
    hdr_cells[1].text = 'MSE'
    hdr_cells[2].text = 'NMSE'
    hdr_cells[3].text = 'PSNR'
    hdr_cells[4].text = 'IF'
    hdr_cells[5].text = 'SSIM'
    hdr_cells[6].text = 'Typ degradacji'
    
    for res in results:
        row_cells = table.add_row().cells
        row_cells[0].text = str(res['params'])
        row_cells[1].text = f"{res['MSE']:.4f}"
        row_cells[2].text = f"{res['NMSE']:.4f}"
        row_cells[3].text = f"{res['PSNR']:.2f}"
        row_cells[4].text = f"{res['IF']:.4f}"
        row_cells[5].text = f"{res['SSIM']:.4f}"
        row_cells[6].text = res['method']

def addImToDoc(document, img, width=6.0):
    buf = BytesIO()
    img_bgr = cv2.cvtColor(img, cv2.COLOR_RGB2BGR)
    _, img_encoded = cv2.imencode('.png', img_bgr)
    buf.write(img_encoded)
    buf.seek(0)
    document.add_picture(buf, width=Inches(width))
    buf.close()

path = pathlib.Path().absolute() / 'l10'
images = ['z1.png', 'z2.png', 'z3.png']
blurs = [
    (cv2.bilateralFilter, {'d': 9, 'sigmaColor': 75, 'sigmaSpace': 75}),
    (cv2.bilateralFilter, {'d': 15, 'sigmaColor': 100, 'sigmaSpace': 100}),
    (cv2.bilateralFilter, {'d': 21, 'sigmaColor': 150, 'sigmaSpace': 150}),
    (cv2.bilateralFilter, {'d': 27, 'sigmaColor': 200, 'sigmaSpace': 200}),
    (cv2.bilateralFilter, {'d': 33, 'sigmaColor': 250, 'sigmaSpace': 250})]

noises =[(szum, {'alpha': 0.1, 'Noise_range': [-25, 25]}),
         (szum, {'alpha': 0.1, 'Noise_range': [-50, 50]}),
         (szum, {'alpha': 0.5, 'Noise_range': [-25, 25]}),
         (szum, {'alpha': 0.5, 'Noise_range': [-50, 50]}),
         (szum, {'alpha': 1, 'Noise_range': [-100, 100]})]

JPEGs = [(JPEGCompress, {'quality': 75}),
         (JPEGCompress, {'quality': 50}),
         (JPEGCompress, {'quality': 35}),
         (JPEGCompress, {'quality': 25}),
         (JPEGCompress, {'quality': 10})]

iterations = 5

img1=cv2.imread(path / 'z1.png')
img1= cv2.cvtColor(img1, cv2.COLOR_BGR2RGB)

img2=cv2.imread(path / 'z2.png')
img2= cv2.cvtColor(img2, cv2.COLOR_BGR2RGB)

img3=cv2.imread(path / 'z3.png')
img3= cv2.cvtColor(img3, cv2.COLOR_BGR2RGB)


document = Document()
document.add_heading('Sprawozdanie 10 Jacewicz Jacek', level=1)
document.add_heading('Zdjęcie nr 1 - Rozmycie', level=2)
addImToDoc(document, img1, width=6.0)

results_blur = []
for i, (blur_func, params) in tqdm(enumerate(blurs),desc='Processing blurs', total=len(blurs)):
    blurred = blur_func(img1, **params)

    document.add_paragraph('Parametry rozmycia: '+str(params))
    plotSM(img1, blurred, title1='Original Image', title2='Blurred Image', document=document)
    results_blur.append({
        'method': 'blur',
        'params': params,
        'MSE': mse(img1, blurred),
        'NMSE': nmse(img1, blurred),
        'PSNR': psnr(img1, blurred),
        'IF': image_fidelity(img1, blurred),
        'SSIM': compute_ssim(img1, blurred)
    })
add_results_table(document, 'Rozmycie (Blur)', results_blur)
document.add_page_break()
document.add_heading('Zdjęcie nr 2 - Szum', level=2)
addImToDoc(document, img2, width=6.0)
results_noise = []
for i, (noise_func, params) in tqdm(enumerate(noises),desc='Processing noises', total=len(noises)):
    noisy = noise_func(img2, **params)
    document.add_paragraph('Parametry szumu: '+str(params))

    plotSM(img2, noisy, title1='Original Image', title2='Noisy Image', document=document)
    results_noise.append({
        'method': 'noise',
        'params': params,
        'MSE': mse(img2, noisy),
        'NMSE': nmse(img2, noisy),
        'PSNR': psnr(img2, noisy),
        'IF': image_fidelity(img2, noisy),
        'SSIM': compute_ssim(img2, noisy)
    })
add_results_table(document, 'Zaszumienie (Noise)', results_noise)
document.add_page_break()
document.add_heading('Zdjęcie nr 3 - kompresja JPG', level=2)
addImToDoc(document, img3, width=6.0)
results_jpeg = []
for i, (jpeg_func, params) in tqdm(enumerate(JPEGs), desc='Processing JPEG compressions', total=len(JPEGs)):
    jpeg_img = jpeg_func(img3, **params)
    document.add_paragraph('Parametry kompresji: '+str(params))
    plotSM(img3, jpeg_img, title1='Original Image', title2='JPEG Compressed Image', document=document)
    results_jpeg.append({
        'method': 'jpeg',
        'params': params,
        'MSE': mse(img3, jpeg_img),
        'NMSE': nmse(img3, jpeg_img),
        'PSNR': psnr(img3, jpeg_img),
        'IF': image_fidelity(img3, jpeg_img),
        'SSIM': compute_ssim(img3, jpeg_img)
    })

add_results_table(document, 'Kompresja JPEG', results_jpeg)
document.add_page_break()

document.add_heading('Analiza korelacji', level=2)
document.add_paragraph('Analiza korelacji dla rozmycia:')

document.save("Sprawozdanie10.docx")